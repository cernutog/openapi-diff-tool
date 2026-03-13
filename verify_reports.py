
from docx import Document
import sys

def verify_files(filenames, search_terms):
    for filename in filenames:
        print(f"--- Verifying {filename} ---")
        doc = Document(filename)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    full_text.append(cell.text)
        
        content = "\n".join(full_text)
        for term in search_terms:
            if term in content:
                print(f"FOUND: '{term}'")
            else:
                print(f"MISSING: '{term}'")

if __name__ == "__main__":
    verify_files(['test_analytic.docx', 'test_impact.docx'], [
        "New brand description of User schema",
        "The user's name (newly updated)",
        "Description changed"
    ])
