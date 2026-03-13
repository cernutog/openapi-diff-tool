import datetime
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import difflib
from dependency_tracer import DependencyTracer

# OXML Helpers
def get_or_add_child(parent, tag_name, ordering=None):
    if ordering is None:
        ordering = []
    child = parent.find(qn(tag_name))
    if child is None:
        insert_index = len(parent)
        if ordering:
            try:
                tag_pos = ordering.index(tag_name)
                for i, existing_child in enumerate(parent):
                    existing_tag = existing_child.tag.split('}')[-1]
                    if f'w:{existing_tag}' in ordering:
                        if ordering.index(f'w:{existing_tag}') > tag_pos:
                            insert_index = i
                            break
            except ValueError:
                pass
        child = OxmlElement(tag_name)
        parent.insert(insert_index, child)
    return child

TBL_PR_ORDER = ['w:tblStyle', 'w:tblpPr', 'w:tblOverlap', 'w:bidiVisual', 'w:tblStyleRowBandSize', 'w:tblStyleColBandSize', 'w:tblW', 'w:jc', 'w:tblCellSpacing', 'w:tblInd', 'w:tblBorders', 'w:shd', 'w:tblLayout', 'w:tblCellMar', 'w:tblLook']
TC_PR_ORDER = ['w:tcW', 'w:gridSpan', 'w:hMerge', 'w:vMerge', 'w:tcBorders', 'w:shd', 'w:noWrap', 'w:tcMar', 'w:textDirection', 'w:tcFitText', 'w:vAlign', 'w:hideMark']

class ImpactDocxGenerator:
    def __init__(self, old_spec, new_spec, diff, old_path=None, new_path=None, variables=None, template_path=None):
        self.old_spec = old_spec
        self.new_spec = new_spec
        self.diff = diff
        self.old_path = old_path
        self.new_path = new_path
        self.variables = variables or {}
        
        # Template Loading Logic
        # 1. Specific template passed in (e.g., template_impact.docx)
        # 2. Fallback to generic 'template.docx'
        # 3. No template
        
        self.template_path = None
        self.has_template = False
        
        if template_path and os.path.exists(template_path):
            self.template_path = template_path
        elif os.path.exists("template.docx"):
            self.template_path = "template.docx"
            
        if self.template_path:
            self.doc = Document(self.template_path)
            self.has_template = True
        else:
            self.doc = Document()
            self.has_template = False
            
        self._setup_styles()
        
        if not self.has_template:
            self._setup_page_layout()
        
        self.analysis_insights = []
        self.checklist_items = []
        
        # Initialize Dependency Tracers
        self.tracer = DependencyTracer(new_spec)
        self.tracer.resolve_transitive_impact()
        
        self.old_tracer = DependencyTracer(old_spec)
        self.old_tracer.resolve_transitive_impact()
            
        self._run_smart_analysis()

    def generate(self, output_path):
        self.doc.add_heading('OpenAPI Comparison - Impact Report', 0)
        self._add_spec_metadata()
        self._add_migration_notice()
        self._add_endpoint_impact_matrix()
        self._add_detailed_component_analysis()
        self._add_technical_deep_dive()
        self._add_implementation_checklist()
        
        # Variable Substitution (Final Step)
        self._process_template_variables()
        
        self.doc.save(output_path)
        print(f"Impact Report generated at {output_path}")

    def _process_template_variables(self):
        """
        Replaces {{ variable }} placeholders in the document with values.
        Prioritizes:
        1. Dynamic Variables (date, time, filenames)
        2. User Static Variables (from Preferences)
        """
        import datetime
        import getpass
        import sys
        
        # 1. Prepare Variables
        context = self.variables.copy()
        
        # Dynamic Defaults
        now = datetime.datetime.now()
        context['date'] = now.strftime('%Y-%m-%d')
        context['time'] = now.strftime('%H:%M')
        context['datetime'] = now.strftime('%Y-%m-%d %H:%M:%S')
        context['original_spec'] = os.path.basename(self.old_path) if self.old_path else "N/A"
        context['new_spec'] = os.path.basename(self.new_path) if self.new_path else "N/A"
        
        # Enriched Variables
        try:
            context['user'] = getpass.getuser()
        except:
            context['user'] = "Unknown"
            
        context['platform'] = sys.platform
        context['tool_version'] = "1.0.0"
        
        def get_size(path):
            if path and os.path.exists(path):
                size_bytes = os.path.getsize(path)
                if size_bytes < 1024: return f"{size_bytes} B"
                elif size_bytes < 1024*1024: return f"{size_bytes/1024:.1f} KB"
                else: return f"{size_bytes/(1024*1024):.1f} MB"
            return "N/A"
            
        context['file_size_old'] = get_size(self.old_path)
        context['file_size_new'] = get_size(self.new_path)
        
        # 2. Iterate and Replace
        # Body Paragraphs
        for p in self.doc.paragraphs:
            self._replace_text_in_paragraph(p, context)
            
        # Tables
        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        self._replace_text_in_paragraph(p, context)
                        
        # Headers & Footers
        for section in self.doc.sections:
            # Header
            for p in section.header.paragraphs:
                self._replace_text_in_paragraph(p, context)
            for table in section.header.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            self._replace_text_in_paragraph(p, context)
                            
            # Footer
            for p in section.footer.paragraphs:
                self._replace_text_in_paragraph(p, context)
            for table in section.footer.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            self._replace_text_in_paragraph(p, context)

    def _replace_text_in_paragraph(self, paragraph, context):
        if '{{' not in paragraph.text:
            return
            
        # Improved approach: Check each run.
        for run in paragraph.runs:
            for key, value in context.items():
                placeholder = f"{{{{ {key} }}}}"
                placeholder_tight = f"{{{{{key}}}}}"
                
                if placeholder in run.text:
                    run.text = run.text.replace(placeholder, str(value))
                if placeholder_tight in run.text:
                    run.text = run.text.replace(placeholder_tight, str(value))


    def _setup_styles(self):
        # Serif Title
        if 'Title' in self.doc.styles:
            style = self.doc.styles['Title']
            style.font.name = 'Georgia'
            style.font.size = Pt(26)
            style.font.bold = True
            style.font.color.rgb = RGBColor(31, 78, 121) # Dark Blue

        # Normal text
        normal = self.doc.styles['Normal']
        normal.font.name = 'Segoe UI'
        normal.font.size = Pt(10)
        normal.paragraph_format.left_indent = Pt(0) # Force reset
        normal.paragraph_format.first_line_indent = Pt(0) # Force reset
        
        # Table Header (Clean - No Background)
        if 'Table Header' not in self.doc.styles:
            s = self.doc.styles.add_style('Table Header', WD_STYLE_TYPE.PARAGRAPH)
            s.base_style = normal
            s.font.bold = True
            s.font.color.rgb = RGBColor(80, 80, 80) # Dark Grey
            s.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            s.paragraph_format.space_before = Pt(6)
            s.paragraph_format.space_after = Pt(6)

        # Table Text
        if 'Table Text' not in self.doc.styles:
            s = self.doc.styles.add_style('Table Text', WD_STYLE_TYPE.PARAGRAPH)
            s.base_style = normal
            s.font.size = Pt(9)
            s.paragraph_format.space_before = Pt(4)
            s.paragraph_format.space_after = Pt(4)

        # Insight Box Title
        if 'Insight Title' not in self.doc.styles:
            s = self.doc.styles.add_style('Insight Title', WD_STYLE_TYPE.PARAGRAPH)
            s.base_style = normal
            s.font.size = Pt(11)
            s.font.bold = True
            s.font.color.rgb = RGBColor(31, 78, 121)
            s.paragraph_format.space_before = Pt(6)
            s.paragraph_format.space_after = Pt(2)

    def _setup_page_layout(self):
        section = self.doc.sections[0]
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)

    def _add_spec_metadata(self):
        # Create a table for Spec Details
        widths = [Inches(1.5), Inches(2.75), Inches(2.75)]
        table = self.doc.add_table(rows=1, cols=3)
        self._remove_all_borders(table)
        self._set_table_fixed_width(table, 7.0)
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        
        # Custom Header Styling (Dark Blue Background, White Text)
        headers = ['Detail', 'Old Specification', 'New Specification']
        row = table.rows[0]
        for i, text in enumerate(headers):
            cell = row.cells[i]
            cell.text = text
            p = cell.paragraphs[0]
            p.style = 'Table Header'
            p.runs[0].font.color.rgb = RGBColor(255, 255, 255) # White
            
            # Dark Blue Background
            tcPr = cell._tc.get_or_add_tcPr()
            shd = get_or_add_child(tcPr, 'w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:fill'), '1F4E79') # Dark Blue
            
            # Borders
            tcBorders = get_or_add_child(tcPr, 'w:tcBorders')
            bottom = get_or_add_child(tcBorders, 'w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '12')
            bottom.set(qn('w:color'), 'FFFFFF')
        
        # Helper to get info
        def get_info(spec, path):
            info = spec.get('info', {})
            return {
                'file': os.path.basename(path) if path else "N/A",
                'title': info.get('title', 'N/A'),
                'version': info.get('version', 'N/A')
            }
            
        old_info = get_info(self.old_spec, self.old_path)
        new_info = get_info(self.new_spec, self.new_path)
        
        # Rows
        rows = [
            ("File Name", old_info['file'], new_info['file']),
            ("API Title", old_info['title'], new_info['title']),
            ("Version", old_info['version'], new_info['version'])
        ]
        
        for label, old_val, new_val in rows:
            row = table.add_row()
            
            # Horizontal Border Only (Light Grey)
            for cell in row.cells:
                tcPr = cell._tc.get_or_add_tcPr()
                tcBorders = get_or_add_child(tcPr, 'w:tcBorders', TC_PR_ORDER)
                for side in ['top', 'left', 'right']:
                    tag = get_or_add_child(tcBorders, f'w:{side}')
                    tag.set(qn('w:val'), 'nil')
                bottom = get_or_add_child(tcBorders, 'w:bottom')
                bottom.set(qn('w:val'), 'single')
                bottom.set(qn('w:sz'), '4')
                bottom.set(qn('w:color'), 'E0E0E0')
            
            row.cells[0].text = label
            row.cells[1].text = str(old_val)
            row.cells[2].text = str(new_val)
            
            # Style Label
            row.cells[0].paragraphs[0].style = 'Table Text'
            row.cells[0].paragraphs[0].runs[0].font.bold = True
            
            # Style Values
            row.cells[1].paragraphs[0].style = 'Table Text'
            row.cells[2].paragraphs[0].style = 'Table Text'
                
        self.doc.add_paragraph().paragraph_format.space_after = Pt(12)

    def _add_migration_notice(self):
        if not self.analysis_insights:
            return

        # Determine Max Severity
        severities = [i.get('severity') for i in self.analysis_insights]
        if 'CRITICAL' in severities:
            max_severity = 'CRITICAL'
            bg_color = 'F8D7DA' # Light Red
            accent_color = '721C24' # Dark Red
            text_color = RGBColor(114, 28, 36)
            intro_text = "Critical breaking changes detected:"
        elif 'HIGH' in severities:
            max_severity = 'HIGH'
            bg_color = 'FFF3CD' # Light Yellow
            accent_color = '856404' # Dark Yellow
            text_color = RGBColor(133, 100, 4)
            intro_text = "Important changes detected:"
        else:
            max_severity = 'LOW'
            bg_color = 'D1ECF1' # Pastel Blue
            accent_color = '0C5460' # Dark Teal
            text_color = RGBColor(12, 84, 96)
            intro_text = "Minor changes detected:"

        # Calculate stats and collect items
        categories = {} # cat_name -> {'count': int, 'items': set, 'severity': str}
        # Show all Critical and High changes in summary
        significant_severities = ['CRITICAL', 'HIGH']
        
        for i in self.analysis_insights:
            sev = i.get('severity')
            if sev in significant_severities:
                cat = i.get('title', 'Unknown')
                if cat not in categories:
                    categories[cat] = {'count': 0, 'items': set(), 'severity': sev}
                categories[cat]['count'] += 1
                if i.get('affected_items'):
                    for item in i['affected_items']:
                        if item: categories[cat]['items'].add(str(item))
        
        # Title OUTSIDE the box
        p_title = self.doc.add_paragraph()
        p_title.paragraph_format.space_before = Pt(12)
        p_title.paragraph_format.space_after = Pt(4)
        run_title = p_title.add_run("MIGRATION IMPACT OVERVIEW")
        run_title.bold = True
        run_title.font.name = 'Arial'
        run_title.font.size = Pt(11)
        run_title.font.color.rgb = RGBColor(80, 80, 80)

        # The Box (Table with Left Accent)
        # 2 Columns: Accent Bar (0.15") | Content (6.85") = Total 7.0"
        tbl = self.doc.add_table(rows=1, cols=2)
        self._remove_all_borders(tbl)
        self._set_table_fixed_width(tbl, 7.0)
        tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
        
        # Column 1: Accent Bar
        cell_accent = tbl.cell(0, 0)
        cell_accent.width = Inches(0.15)
        tcPr = cell_accent._tc.get_or_add_tcPr()
        shd = get_or_add_child(tcPr, 'w:shd', TC_PR_ORDER)
        shd.set(qn('w:fill'), accent_color)
        
        # Column 2: Content
        cell_content = tbl.cell(0, 1)
        cell_content.width = Inches(6.85)
        tcPr = cell_content._tc.get_or_add_tcPr()
        shd = get_or_add_child(tcPr, 'w:shd', TC_PR_ORDER)
        shd.set(qn('w:fill'), bg_color)
        
        # Add Padding to Content Cell
        tcMar = get_or_add_child(tcPr, 'w:tcMar', TC_PR_ORDER)
        for side in ['top', 'bottom', 'left', 'right']:
             mar = get_or_add_child(tcMar, f'w:{side}')
             mar.set(qn('w:w'), '120') # ~6pt padding
             mar.set(qn('w:type'), 'dxa')

        # Content inside box
        p = cell_content.paragraphs[0]
        
        # Intro
        run = p.add_run(intro_text)
        run.font.color.rgb = text_color
        run.font.size = Pt(10)
        run.bold = True
        p.add_run("\n")

        # Bullet points
        for cat, data in categories.items():
            count = data['count']
            items = sorted(list(data['items']))
            
            # Simple Pluralization
            display_cat = cat
            if count > 1:
                if "Removed" in cat:
                     display_cat = cat.replace("Removed", "Removals")
                     if "Property" in cat: display_cat = "Property Removals"
                     if "Endpoint" in cat: display_cat = "Endpoint Removals"
                     if "Parameter" in cat: display_cat = "Parameter Removals"
                elif cat.endswith('y'): display_cat = cat[:-1] + "ies"
                elif not cat.endswith('s'): display_cat = cat + "s"
            
            main_text = f" \u2022 {count} {display_cat}"
            if items:
                # Limit to first 10 items for better detail in summary
                limit = 10
                item_str = ", ".join(items[:limit])
                if len(items) > limit:
                    item_str += f", and {len(items)-limit} more"
                main_text += f": {item_str}" # Changed (items) to : items for more prominence
            
            run_bullet = p.add_run(main_text)
            run_bullet.font.color.rgb = text_color
            run_bullet.font.size = Pt(10)
            p.add_run("\n")
        
        # Footer message (Refined Tone)
        if max_severity == 'CRITICAL':
            msg = "Backward compatibility is compromised. A migration plan is required."
        elif max_severity == 'HIGH':
            msg = "Significant changes to contract. Client updates likely required."
        else:
            msg = "Standard regression testing recommended."

        p.add_run("\n") # Extra spacing before footer
        run2 = p.add_run(msg)
        run2.font.color.rgb = text_color
        run2.font.size = Pt(9)
        run2.italic = True
        
        self.doc.add_paragraph()

    def _add_section_header(self, number, title):
        # Use Heading 1 for Navigation Pane support
        p = self.doc.add_heading(f"{number}. {title}", level=1)
        
        # Apply custom styling on top of Heading 1
        p.paragraph_format.space_before = Pt(24)
        p.paragraph_format.space_after = Pt(12)
        
        # Reset any indent from Heading 1 if needed
        p.paragraph_format.left_indent = Pt(0)
        
        pPr = p._p.get_or_add_pPr()
        pbdr = get_or_add_child(pPr, 'w:pBdr')
        left = get_or_add_child(pbdr, 'w:left')
        left.set(qn('w:val'), 'single')
        left.set(qn('w:sz'), '24')
        left.set(qn('w:space'), '12')
        left.set(qn('w:color'), '4A90E2') # Blue
        
        for run in p.runs:
            run.font.name = 'Arial'
            run.font.size = Pt(11)
            run.font.all_caps = True
            run.font.color.rgb = RGBColor(0, 0, 0) # Force black or keep default? Let's use Black.
        
        # Add an explicit spacer paragraph to ensure separation
        spacer = self.doc.add_paragraph()
        spacer.paragraph_format.space_after = Pt(4)
        spacer.paragraph_format.line_spacing = Pt(4)

    def _remove_all_borders(self, table):
        """
        Explicitly removes all borders from the table properties.
        """
        tbl = table._tbl
        tblPr = tbl.tblPr
        tblBorders = get_or_add_child(tblPr, 'w:tblBorders', TBL_PR_ORDER)
        
        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = get_or_add_child(tblBorders, f'w:{border_name}')
            border.set(qn('w:val'), 'nil')

    def _set_table_fixed_width(self, table, width_inches):
        """
        Forces the table to a fixed width using OXML, overriding autofit.
        """
        tbl = table._tbl
        tblPr = tbl.tblPr
        
        # 1. Set Table Width (w:tblW) to fixed value
        tblW = get_or_add_child(tblPr, 'w:tblW', TBL_PR_ORDER)
        tblW.set(qn('w:w'), str(int(width_inches * 1440))) # 1440 twips per inch
        tblW.set(qn('w:type'), 'dxa')
        
        # 2. Set Table Layout (w:tblLayout) to fixed
        tblLayout = get_or_add_child(tblPr, 'w:tblLayout', TBL_PR_ORDER)
        tblLayout.set(qn('w:type'), 'fixed')

    def _add_endpoint_impact_matrix(self):
        self._add_section_header("1", "ENDPOINT IMPACT MATRIX")
        
        table = self.doc.add_table(rows=1, cols=3)
        self._remove_all_borders(table)
        self._set_table_fixed_width(table, 7.0) # Increased to 7.0 inches (Max Alignment)
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        
        # Optimized Column Widths (Total 7.0)
        # OLD: Method 0.5, Impact 4.7, Resource 1.8
        # NEW: Method 0.8, Impact 4.4, Resource 1.8
        widths = [Inches(1.8), Inches(0.8), Inches(4.4)]
        for i, width in enumerate(widths):
            table.cell(0, i).width = width
            
        # Clean Header
        hdr_cells = table.rows[0].cells
        headers = ["ENDPOINT RESOURCE", "METHOD", "DETAILED TECHNICAL IMPACT"]
        for i, text in enumerate(headers):
            hdr_cells[i].text = text
            p = hdr_cells[i].paragraphs[0]
            p.style = 'Table Header'
            # Bottom border only for header
            tcPr = hdr_cells[i]._tc.get_or_add_tcPr()
            tcBorders = get_or_add_child(tcPr, 'w:tcBorders', TC_PR_ORDER)
            bottom = get_or_add_child(tcBorders, 'w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '12')
            bottom.set(qn('w:color'), '000000')

        # 0. Collect all relevant endpoints
        endpoints_to_show = set() # (path, method, source_data)
        
        # Modified Paths (Structural)
        if hasattr(self.diff, 'modified_paths'):
            for path, p_changes in self.diff.modified_paths.items():
                if 'modified_ops' in p_changes:
                    for op, op_changes in p_changes['modified_ops'].items():
                        endpoints_to_show.add((path, op.upper()))
                if 'removed_ops' in p_changes:
                    for op in p_changes['removed_ops']:
                        endpoints_to_show.add((path, op.upper()))

        # New Paths
        if hasattr(self.diff, 'new_paths'):
            for path in self.diff.new_paths:
                p_item = self.new_spec.get('paths', {}).get(path, {})
                methods = [m.upper() for m in p_item.keys() if m.lower() in ['get', 'post', 'put', 'delete', 'patch', 'options', 'head', 'trace']]
                for m in (methods if methods else ['GET']):
                    endpoints_to_show.add((path, m))

        # Removed Paths
        if hasattr(self.diff, 'removed_paths'):
            for path in self.diff.removed_paths:
                p_item = self.old_spec.get('paths', {}).get(path, {})
                methods = [m.upper() for m in p_item.keys() if m.lower() in ['get', 'post', 'put', 'delete', 'patch', 'options', 'head', 'trace']]
                for m in (methods if methods else ['GET']):
                    endpoints_to_show.add((path, m))

        # Indirectly Impacted Endpoints (from tracer/insights)
        for method, path in self.endpoint_insights.keys():
            endpoints_to_show.add((path, method))

        # Sort endpoints by path then method
        sorted_endpoints = sorted(list(endpoints_to_show))

        # Add Rows
        for path, method in sorted_endpoints:
            # Reconstruct 'changes' if available in diff
            p_changes = self.diff.modified_paths.get(path, {}) if hasattr(self.diff, 'modified_paths') else {}
            op_changes = {}
            if 'modified_ops' in p_changes and method.lower() in p_changes['modified_ops']:
                op_changes = p_changes['modified_ops'][method.lower()]
            elif 'removed_ops' in p_changes and method.lower() in p_changes['removed_ops']:
                op_changes = {'removed': True}
            
            # Check for new/removed path
            if hasattr(self.diff, 'new_paths') and path in self.diff.new_paths:
                op_changes = {'new': True}
            elif hasattr(self.diff, 'removed_paths') and path in self.diff.removed_paths:
                op_changes = {'removed': True}

            self._add_impact_row(table, path, method, op_changes)
        
        # Add Spacer after table
        spacer = self.doc.add_paragraph()
        spacer.paragraph_format.space_after = Pt(24)

    def _add_impact_row(self, table, path, method, changes):
        row = table.add_row()
        
        # Horizontal Border Only (Light Grey)
        for cell in row.cells:
            tcPr = cell._tc.get_or_add_tcPr()
            tcBorders = get_or_add_child(tcPr, 'w:tcBorders', TC_PR_ORDER)
            for side in ['top', 'left', 'right']:
                tag = get_or_add_child(tcBorders, f'w:{side}')
                tag.set(qn('w:val'), 'nil')
            bottom = get_or_add_child(tcBorders, 'w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '4')
            bottom.set(qn('w:color'), 'E0E0E0')

        row.cells[0].text = path
        row.cells[0].paragraphs[0].style = 'Table Text'
        row.cells[0].paragraphs[0].runs[0].font.bold = True
        
        row.cells[1].text = method
        row.cells[1].paragraphs[0].style = 'Table Text'
        
        cell = row.cells[2]
        p = cell.paragraphs[0]
        p.style = 'Table Text'
        
        impacts = self._analyze_impact_for_row(path, method, changes)
        for severity, msg in impacts:
            self._add_pill_badge(p, severity)
            p.add_run(f" {msg}\n")

    def _analyze_impact_for_row(self, path, method, changes):
        """
        Calculates impact using deep structural analysis and heuristic insights.
        """
        impacts = []
        
        # 1. Structural Changes (Mandatory Details)
        if changes.get('new'):
            return [('LOW', "New endpoint available.")]
        if changes.get('removed'):
            return [('CRITICAL', "Endpoint removed. Breaking change.")]

        # A) HEURISTIC INSIGHTS (Advice/Context)
        methods = [m.strip() for m in method.split(",")]
        related_insights = []
        for m in methods:
            related_insights.extend(self.endpoint_insights.get((m.upper(), path), []))
        
        # Add high-level insights first
        added_titles = set()
        for i in related_insights:
            msg = i['title']
            if i.get('affected_items'):
                msg += f": {', '.join(map(str, i['affected_items']))}"
            
            if msg not in added_titles:
                impacts.append((i['severity'], msg))
                added_titles.add(msg)

        # B) DEEP STRUCTURAL SCAN (Technical Truth)
        # This recursively finds all removals directly in the diff for this operation
        details = self._crawl_diff_for_impacts(changes)
        for detail in details:
            impacts.append(('CRITICAL', detail))

        # 3. Generic Fallbacks for things not caught above
        if not impacts:
            if 'responses' in changes:
                impacts.append(('HIGH', "Response schema updated."))
            else:
                impacts.append(('INFO', "Documentation or metadata updated."))
                
        return impacts

    def _crawl_diff_for_impacts(self, diff_node, prefix=""):
        """
        Recursively extracts removed items from a diff structure.
        """
        found = []
        if not isinstance(diff_node, dict): return found

        # 1. Handle Removed Parameters
        if 'parameters' in diff_node and diff_node['parameters'].get('removed'):
            p_list = ", ".join(diff_node['parameters']['removed'])
            found.append(f"Parameter Removed: {p_list}")

        # 2. Handle Removed Properties
        if 'properties' in diff_node and diff_node['properties'].get('removed'):
            p_list = ", ".join(diff_node['properties']['removed'])
            loc = f" ({prefix})" if prefix else ""
            found.append(f"Property Removed{loc}: {p_list}")

        # 3. Handle Enum removals
        if 'enum' in diff_node and diff_node['enum'].get('removed'):
             e_list = ", ".join(map(str, diff_node['enum']['removed']))
             found.append(f"Enum Values Removed: {e_list}")

        # 4. Handle Specific Logic (Body/Responses/Items)
        # Scan content maps (RequestBody/Responses)
        if 'content' in diff_node and 'modified' in diff_node.get('content', {}):
            for mt, mt_diff in diff_node['content']['modified'].items():
                found.extend(self._crawl_diff_for_impacts(mt_diff, f"Body {mt}"))

        if 'responses' in diff_node and 'modified' in diff_node.get('responses', {}):
            for code, code_diff in diff_node['responses']['modified'].items():
                found.extend(self._crawl_diff_for_impacts(code_diff, f"Resp {code}"))

        if 'schema' in diff_node:
            found.extend(self._crawl_diff_for_impacts(diff_node['schema'], prefix))

        if 'items' in diff_node:
             found.extend(self._crawl_diff_for_impacts(diff_node['items'], f"{prefix}[]"))

        # Recurse into modified properties to find nested removals
        if 'properties' in diff_node and 'modified' in diff_node.get('properties', {}):
            for p_name, p_diff in diff_node['properties']['modified'].items():
                new_prefix = f"{prefix}.{p_name}" if prefix else p_name
                found.extend(self._crawl_diff_for_impacts(p_diff, new_prefix))

        return list(dict.fromkeys(found)) # Deduplicate
            
        # Prioritize CRITICAL over others if multiple
        impacts.sort(key=lambda x: {'CRITICAL': 0, 'HIGH': 1, 'MEDIUM': 2, 'LOW': 3, 'INFO': 4}.get(x[0], 5))
        
        return impacts[:3] # Limit to top 3 for space

    def _add_detailed_component_analysis(self):
        self._add_section_header("2", "DETAILED COMPONENT ANALYSIS")
        
        # Gather all items (Modified + Pure Renames) for all types
        comp_types = ['schemas', 'parameters', 'responses', 'headers', 'securitySchemes', 'examples', 'links', 'callbacks']
        items_to_show = []
        
        for c_type in comp_types:
            # 1. Modified
            if hasattr(self.diff, 'modified_components') and c_type in self.diff.modified_components:
                for s_name, s_changes in self.diff.modified_components[c_type].items():
                    items_to_show.append({'name': s_name, 'data': s_changes, 'type': 'modified', 'c_type': c_type})

            # 2. Pure Renames
            if hasattr(self.diff, 'renamed_components') and c_type in self.diff.renamed_components:
                mod_keys = self.diff.modified_components.get(c_type, {}).keys()
                for old, new in self.diff.renamed_components[c_type].items():
                    if old not in mod_keys:
                        items_to_show.append({'name': old, 'data': {'new_name': new}, 'type': 'renamed', 'c_type': c_type})

            # 3. Removed
            if hasattr(self.diff, 'removed_components') and c_type in self.diff.removed_components:
                for name in self.diff.removed_components[c_type]:
                    # Ignore if it was part of a rename
                    if name not in self.diff.renamed_components.get(c_type, {}):
                        items_to_show.append({'name': name, 'data': {'removed': True}, 'type': 'removed', 'c_type': c_type})
        
        if items_to_show:
            table = self.doc.add_table(rows=1, cols=4)
            self._remove_all_borders(table)
            self._set_table_fixed_width(table, 7.0) 
            table.alignment = WD_TABLE_ALIGNMENT.LEFT
            
            widths = [Inches(1.5), Inches(0.8), Inches(3.3), Inches(1.4)]
            for i, width in enumerate(widths):
                table.cell(0, i).width = width
            
            hdr_cells = table.rows[0].cells
            headers = ["COMPONENT", "TYPE", "CHANGE DETAILS", "AFFECTED ENDPOINTS"]
            for i, text in enumerate(headers):
                hdr_cells[i].text = text
                p = hdr_cells[i].paragraphs[0]
                p.style = 'Table Header'
                tcPr = hdr_cells[i]._tc.get_or_add_tcPr()
                tcBorders = get_or_add_child(tcPr, 'w:tcBorders', TC_PR_ORDER)
                bottom = get_or_add_child(tcBorders, 'w:bottom')
                bottom.set(qn('w:val'), 'single')
                bottom.set(qn('w:sz'), '12')
                bottom.set(qn('w:color'), '000000')

            # Sort by component type then name
            items_to_show.sort(key=lambda x: (x['c_type'], x['name']))

            for item in items_to_show:
                self._add_component_row(table, item)

        # Add Spacer after table
        spacer = self.doc.add_paragraph()
        spacer.paragraph_format.space_after = Pt(24)

    def _add_component_row(self, table, item):
        s_name = item['name']
        s_changes = item['data']
        item_type = item['type']
        c_type = item.get('c_type', 'schemas')

        row = table.add_row()
        
        # Horizontal Border Only (Light Grey)
        for cell in row.cells:
            tcPr = cell._tc.get_or_add_tcPr()
            tcBorders = get_or_add_child(tcPr, 'w:tcBorders', TC_PR_ORDER)
            for side in ['top', 'left', 'right']:
                tag = get_or_add_child(tcBorders, f'w:{side}')
                tag.set(qn('w:val'), 'nil')
            bottom = get_or_add_child(tcBorders, 'w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '4')
            bottom.set(qn('w:color'), 'E0E0E0')

        # Column 1: Name (Handle Rename)
        display_name = s_name
        rename_note = ""
        renamed_map = self.diff.renamed_components.get(c_type, {})
        
        if item_type == 'renamed':
            display_name = f"{s_name} \u2192 {s_changes['new_name']}"
        elif item_type == 'modified':
            new_name = renamed_map.get(s_name)
            if new_name:
                display_name = new_name
                rename_note = f"\n(was {s_name})"

        row.cells[0].text = display_name
        row.cells[0].paragraphs[0].style = 'Table Text'
        row.cells[0].paragraphs[0].runs[0].font.bold = True
        if rename_note:
            run = row.cells[0].paragraphs[0].add_run(rename_note)
            run.font.size = Pt(8)
            run.font.italic = True
            run.font.color.rgb = RGBColor(100, 100, 100)
        
        # Column 2: Type
        row.cells[1].text = c_type[:-1].capitalize() if not c_type.endswith('ies') else c_type[:-3].capitalize() + 'y'
        row.cells[1].paragraphs[0].style = 'Table Text'
        
        # Column 3: Details
        cell = row.cells[2]
        p = cell.paragraphs[0]
        p.style = 'Table Text'
        
        if item_type == 'renamed':
            self._add_pill_badge(p, "RENAMED")
            p.add_run(f" Component renamed. Content is identical.")
        elif item_type == 'removed':
            row.cells[2].text = f"The entire {c_type[:-1]} definition has been removed. Review dependencies."
            row.cells[2].paragraphs[0].style = 'Table Text'
            row.cells[2].paragraphs[0].runs[0].font.color.rgb = RGBColor(180, 0, 0) # Red-ish
        else: # modified or renamed with changes
            if rename_note:
                 self._add_pill_badge(p, "RENAMED")
            if c_type == 'schemas':
                self._render_schema_diff_details(p, s_changes)
            else:
                self._render_general_diff_details(p, s_changes)
        display_name_clean = s_name
        if item_type == 'renamed':
             display_name_clean = s_changes['new_name']
        elif item_type == 'modified':
             if s_name in renamed_map:
                 display_name_clean = renamed_map[s_name]

        impacts = self.tracer.get_impacted_endpoints(display_name_clean)
        
        row.cells[3].paragraphs[0].style = 'Table Text'
        if impacts:
            unique_paths = sorted(list(set([f"{i['method']} {i['path']}" for i in impacts])))
            DISPLAY_LIMIT = 5
            for p_str in unique_paths[:DISPLAY_LIMIT]:
                row.cells[3].add_paragraph(p_str, style='Table Text')
            if len(unique_paths) > DISPLAY_LIMIT:
                row.cells[3].add_paragraph(f"... +{len(unique_paths)-DISPLAY_LIMIT} more", style='Table Text')
        else:
             row.cells[3].text = "-"
             row.cells[3].paragraphs[0].style = 'Table Text'

    def _render_general_diff_details(self, p, changes):
        """
        Renders simple attribute changes for a component in a single paragraph.
        """
        first = True
        for key, val in changes.items():
            if key == '__rename_info__': continue
            if isinstance(val, dict) and 'old' in val and 'new' in val:
                if not first:
                    p.add_run("\n")
                first = False
                
                # Attribute Name
                run = p.add_run(f"{key}: ")
                run.font.bold = True
                
                # Values
                if key == 'description':
                    # Simplified rich diff for impact report (just show old/new inline)
                    p.add_run(f"Changed")
                else:
                    old_v = str(val['old'])
                    new_v = str(val['new'])
                    # Limit length
                    if len(old_v) > 50: old_v = old_v[:47] + "..."
                    if len(new_v) > 50: new_v = new_v[:47] + "..."
                    p.add_run(f"'{old_v}' \u2192 '{new_v}'")

    def _render_schema_diff_details(self, p, changes):
        # 1. Properties
        if 'properties' in changes:
            props = changes['properties']
            if 'new' in props:
                for prop in props['new']:
                    self._add_symbol_run(p, "+", "28A745") # Green
                    p.add_run(f" {prop} (new)\n")
            if 'removed' in props:
                for prop in props['removed']:
                    self._add_symbol_run(p, "-", "DC3545") # Red
                    p.add_run(f" {prop} (removed)\n")
            if 'modified' in props:
                for prop, p_diff in props['modified'].items():
                    self._add_symbol_run(p, "~", "FD7E14") # Orange
                    desc = self._summarize_diff(p_diff)
                    p.add_run(f" {prop}: {desc}\n")

        # 2. Required
        if 'required' in changes:
            old_val = changes['required'].get('old')
            new_val = changes['required'].get('new')
            old_req = set(old_val if old_val is not None else [])
            new_req = set(new_val if new_val is not None else [])
            added = new_req - old_req
            removed = old_req - new_req
            
            for item in added:
                self._add_symbol_run(p, "!", "DC3545") # Red (Breaking)
                p.add_run(f" Required: {item} added\n")
            for item in removed:
                self._add_symbol_run(p, "~", "FD7E14") # Orange
                p.add_run(f" Required: {item} removed\n")

        # Combinators (oneOf, etc.)
        for comb in ['oneOf', 'anyOf', 'allOf']:
            if comb in changes:
                c_diff = changes[comb]
                if isinstance(c_diff, dict) and 'added' in c_diff:
                        for item in c_diff['added']:
                            ref = item.get('$ref', 'Inline Schema')
                            self._add_symbol_run(p, "+", "17A2B8") # Cyan
                            p.add_run(f" {comb}: Added option {ref}\n")
            
        # 4. Description
        if 'description' in changes:
            d_diff = changes['description']
            self._add_symbol_run(p, "~", "FD7E14") # Orange
            p.add_run(" Description changed: ")
            self._render_rich_diff_inline(p, d_diff.get('old', ''), d_diff.get('new', ''))
            p.add_run("\n")

    def _add_technical_deep_dive(self):
        self._add_section_header("3", "TECHNICAL DEEP DIVE")
        
        if not self.analysis_insights:
            self.doc.add_paragraph("No significant technical risks detected.")
            return

        # Static Descriptions for Deep Dive (Value-Add vs Duplication)
        RULE_DESCRIPTIONS = {
            'E01': "Endpoints or HTTP methods have been completely removed. This is a definitive breaking change. Clients attempting to access these resources will receive HTTP 404 or 405 errors.",
            'E03': "The 'operationId' field has been modified. This will break any auto-generated SDKs or client libraries that rely on method names derived from this ID.",
            'E04': "Endpoints have been marked as deprecated. While currently functional, they are slated for removal. Clients should be migrated to alternative endpoints.",
            'P01': "Input parameters have been removed from operations. Clients sending these parameters may receive HTTP 400 errors or have the data silently ignored.",
            'P02': "New required parameters have been added. Existing clients unaware of these parameters will fail input validation (HTTP 400).",
            'P04': "Parameters that were previously optional are now mandatory. This tightens the contract and will break clients that omit these parameters.",
            'P07': "The data type of input parameters has changed. This is a fundamental contract violation that will cause serialization/validation errors.",
            'S01': "Properties have been removed from data schemas. Clients relying on these fields for business logic will break or receive incomplete data.",
            'S02': "New required properties have been added to schemas. If these schemas are used in requests (write), clients must be updated. If used in responses (read), clients must handle the new data.",
            'S03': "Property data types have changed. This breaks JSON deserialization in strongly-typed languages (e.g., Java, C#, Go).",
            'S12': "Polymorphic structures (oneOf/anyOf) have been modified. New subtypes may be returned that legacy clients cannot deserialize.",
            'B05': "Request bodies are now required for operations that previously allowed empty bodies.",
            'B03': "Supported Content-Types have been removed. Clients using these media types will receive HTTP 415 Unsupported Media Type."
        }

        # Aggregate insights by Rule ID
        grouped_insights = {}
        for insight in self.analysis_insights:
            rid = insight.get('rule_id')
            if rid not in grouped_insights:
                grouped_insights[rid] = {
                    'title': insight['title'],
                    'contexts': [],
                    'severity': insight['severity']
                }
            if insight.get('context'):
                grouped_insights[rid]['contexts'].append(insight['context'])

        # Render Aggregated Insights
        for i, (rid, data) in enumerate(grouped_insights.items(), 1):
            p_title = self.doc.add_paragraph()
            p_title.style = 'Insight Title'
            p_title.text = f"3.{i} {data['title']}"
            
            # Use Static Description if available, else fallback to generic
            desc = RULE_DESCRIPTIONS.get(rid, "Technical contract modification detected. Review specific changes in the sections above.")
            p_desc = self.doc.add_paragraph(desc)
            p_desc.style = 'Normal'
            
            if data['contexts']:
                p_ctx = self.doc.add_paragraph()
                p_ctx.paragraph_format.left_indent = Inches(0.25)
                p_ctx.paragraph_format.space_after = Pt(12)
                
                # Limit contexts to avoid massive lists
                ctx_list = data['contexts']
                if len(ctx_list) > 10:
                    display_list = ", ".join(ctx_list[:10]) + f", and {len(ctx_list)-10} others..."
                else:
                    display_list = ", ".join(ctx_list)

                run = p_ctx.add_run("Affected Areas: " + display_list)
                run.font.size = Pt(9)
                run.italic = True
                run.font.color.rgb = RGBColor(80, 80, 80)

    def _add_implementation_checklist(self):
        self._add_section_header("4", "IMPLEMENTATION CHECKLIST")
        
        if not self.checklist_items:
            self.doc.add_paragraph("No specific implementation steps required.")
            return
            
        for item in self.checklist_items:
            p = self.doc.add_paragraph()
            run = p.add_run("\u2610 ")
            run.font.size = Pt(12)
            p.add_run(item)

    def _run_smart_analysis(self):
        from heuristic_engine import HeuristicEngine, Severity
        
        engine = HeuristicEngine(self.diff)
        insights_objects = engine.run()
        
        # Map Insights to Report Format (List of Dicts for internal consistency)
        self.analysis_insights = []
        for insight in insights_objects:
            self.analysis_insights.append({
                'title': insight.title,
                'description': insight.description,
                'severity': insight.severity.value,
                'rule_id': insight.rule_id,
                'context': insight.context,
                'affected_items': insight.affected_items
            })
            
            # Generate Checklist Items based on Rule ID
            if insight.rule_id == 'E01':
                self.checklist_items.append(f"Remove usage of {insight.context}.")
            elif insight.rule_id == 'P01':
                self.checklist_items.append(f"Update calls to {insight.context} to remove deleted parameter.")
            elif insight.rule_id == 'P02':
                self.checklist_items.append(f"Update calls to {insight.context} to include new required parameter.")
            elif insight.rule_id == 'S02':
                self.checklist_items.append(f"Update payload builders for {insight.context} to include new required properties.")
            elif insight.rule_id == 'S12':
                self.checklist_items.append(f"Update deserializers for {insight.context} to handle new polymorphic types.")
            elif insight.rule_id == 'B05':
                self.checklist_items.append(f"Ensure request body is provided for {insight.context}.")

        # Collect insights per endpoint for matrix synchronization
        self.endpoint_insights = {}
        for insight in self.analysis_insights:
            ctx = insight.get('context')
            if not ctx: continue
            
            # ROUTING LOGIC:
            # 1. If it starts with "Component: ", it belongs in Section 2, NOT Section 1 (unless traced)
            if ctx.startswith("Component: "):
                # If it's a schema, we TRACE it to endpoints
                if "Schema: " in ctx:
                    schema_name = ctx.split(": ")[-1]
                    # Use BOTH tracers to find usages (covers both modifications and removals)
                    impacted = self.tracer.get_impacted_endpoints(schema_name)
                    impacted_old = self.old_tracer.get_impacted_endpoints(schema_name)
                    
                    # Merge and deduplicate impacted endpoints
                    merged_impacted = { (i['method'], i['path']) for i in (impacted + impacted_old) }
                    
                    for m, p in merged_impacted:
                        key = (m.upper(), p)
                        self.endpoint_insights.setdefault(key, []).append(insight)
                # Others (Examples, etc.) stay only in component analysis section
                continue
            
            # 2. Traditional Endpoint Mapping (METHOD PATH)
            parts = ctx.split(" ", 1)
            if len(parts) == 2:
                m_upper = parts[0].strip().upper()
                p_val = parts[1].strip()
                # Validation: Does the method look like an HTTP method?
                if m_upper in ['GET', 'POST', 'PUT', 'DELETE', 'PATCH', 'OPTIONS', 'HEAD', 'TRACE']:
                    key = (m_upper, p_val)
                    self.endpoint_insights.setdefault(key, []).append(insight)
        
        # Deduplicate Checklist
        self.checklist_items = list(dict.fromkeys(self.checklist_items))

        # Default if empty
        if not self.analysis_insights:
             self.analysis_insights.append({
                'title': 'General Maintenance',
                'description': "Changes appear to be minor or additive. Standard regression testing is recommended.",
                'severity': 'LOW',
                'rule_id': 'GEN',
                'context': None,
                'affected_items': None
            })

    def _summarize_diff(self, diff):
        parts = []
        if 'type' in diff: parts.append(f"Type {diff['type']['old']}->{diff['type']['new']}")
        if 'enum' in diff: parts.append("Enum values changed")
        if 'pattern' in diff: parts.append("Regex pattern changed")
        if 'minimum' in diff or 'maximum' in diff: parts.append("Range constraints changed")
        if 'description' in diff: parts.append("Description changed")
        return ", ".join(parts) if parts else "Constraint changed"

    def _add_symbol_run(self, paragraph, symbol, color_hex):
        run = paragraph.add_run(symbol)
        run.font.name = 'Consolas'
        run.font.bold = True
        run.font.color.rgb = RGBColor.from_string(color_hex)

    def _add_pill_badge(self, paragraph, text):
        # Add spacing for "air"
        run = paragraph.add_run(f"  {text}  ")
        run.font.size = Pt(7)
        run.font.bold = True
        run.font.name = 'Segoe UI'
        
        rPr = run._r.get_or_add_rPr()
        shd = get_or_add_child(rPr, 'w:shd')
        shd.set(qn('w:val'), 'clear')
        
        # Pastel colors with Dark Text for better readability
        if text == 'CRITICAL': 
            shd.set(qn('w:fill'), 'F8D7DA') # Light Red
            run.font.color.rgb = RGBColor(114, 28, 36) # Dark Red
        elif text == 'HIGH': 
            shd.set(qn('w:fill'), 'FFF3CD') # Light Yellow/Orange
            run.font.color.rgb = RGBColor(133, 100, 4) # Dark Yellow
        elif text == 'RELAXED': 
            shd.set(qn('w:fill'), 'D4EDDA') # Light Green
            run.font.color.rgb = RGBColor(21, 87, 36) # Dark Green
        elif text == 'LOW': # Explicitly handle LOW
            shd.set(qn('w:fill'), 'D1ECF1') # Pastel Blue
            run.font.color.rgb = RGBColor(12, 84, 96) # Dark Blue/Teal
        else: 
            shd.set(qn('w:fill'), 'E2E3E5') # Light Grey
            run.font.color.rgb = RGBColor(56, 61, 65) # Dark Grey

    def _render_rich_diff_inline(self, p, text_old, text_new):
        """Renders description diff inline with line splitting and color accuracy."""
        if not isinstance(text_old, str): text_old = str(text_old or "")
        if not isinstance(text_new, str): text_new = str(text_new or "")
        
        # Split into lines
        lines_old = text_old.splitlines(keepends=True)
        lines_new = text_new.splitlines(keepends=True)
        
        s = difflib.SequenceMatcher(None, lines_old, lines_new, autojunk=False)
        opcodes = s.get_opcodes()
        
        # Structural split: separate modifications from removals/additions
        refined_opcodes = []
        for tag, i1, i2, j1, j2 in opcodes:
            if tag == 'replace':
                old_count = i2 - i1
                new_count = j2 - j1
                if old_count > new_count:
                    refined_opcodes.append(('replace', i1, i1 + new_count, j1, j2))
                    refined_opcodes.append(('delete', i1 + new_count, i2, j2, j2))
                elif new_count > old_count:
                    refined_opcodes.append(('replace', i1, i2, j1, j1 + old_count))
                    refined_opcodes.append(('insert', i2, i2, j1 + old_count, j2))
                else:
                    refined_opcodes.append((tag, i1, i2, j1, j2))
            else:
                refined_opcodes.append((tag, i1, i2, j1, j2))

        def apply_shading(run, color_hex):
            rPr = run._r.get_or_add_rPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:fill'), color_hex)
            rPr.append(shd)

        def render_word_diff_inline(para, t_old, t_new):
            import re
            def split_words(text):
                return re.findall(r'\w+|[^\w\s]|\s+', text)
            
            w_old = split_words(t_old)
            w_new = split_words(t_new)
            ws = difflib.SequenceMatcher(None, w_old, w_new, autojunk=False)
            w_ops = ws.get_opcodes()
            
            # Refined merging for inline
            merged_w_ops = []
            wi = 0
            while wi < len(w_ops):
                tag, i1, i2, j1, j2 = w_ops[wi]
                if tag == 'equal':
                    merged_w_ops.append((tag, i1, i2, j1, j2))
                    wi += 1
                else:
                    sw_i1, lw_i2 = i1, i1
                    sw_j1, lw_j2 = j1, j2
                    # Track what we actually saw
                    actual_o = tag in ('delete', 'replace')
                    actual_n = tag in ('insert', 'replace')
                    sw_i1, lw_i2 = i1, i2
                    sw_j1, lw_j2 = j1, j2
                    
                    wk = wi + 1
                    while wk < len(w_ops):
                        nt, ni1, ni2, nj1, nj2 = w_ops[wk]
                        is_bridge = False
                        if nt == 'equal':
                            eq_text = "".join(w_old[ni1:ni2])
                            if len(eq_text) <= 3 and '\n' not in eq_text: is_bridge = True
                        if nt != 'equal' or is_bridge:
                            if nt in ('delete', 'replace'): actual_o = True
                            if nt in ('insert', 'replace'): actual_n = True
                            lw_i2, lw_j2 = ni2, nj2
                            wk += 1
                        else: break
                    wt = 'replace' if (actual_o and actual_n) else ('delete' if actual_o else 'insert')
                    merged_w_ops.append((wt, sw_i1, lw_i2, sw_j1, lw_j2))
                    wi = wk

            # Inline rendering...
            for wt, wi1, wi2, wj1, wj2 in merged_w_ops:
                txt = "".join(w_old[wi1:wi2])
                if wt == 'equal': para.add_run(txt)
                elif wt == 'delete': apply_shading(para.add_run(txt), "F8D7DA")
                elif wt == 'replace': apply_shading(para.add_run(txt), "FFF3CD")
            
            para.add_run(" \u2192 ")
            
            for wt, wi1, wi2, wj1, wj2 in merged_w_ops:
                txt = "".join(w_new[wj1:wj2])
                if wt == 'equal': para.add_run(txt)
                elif wt == 'insert': apply_shading(para.add_run(txt), "D4EDDA")
                elif wt == 'replace': apply_shading(para.add_run(txt), "FFF3CD")

        # Process in sequence
        for tag, i1, i2, j1, j2 in refined_opcodes:
            txt_o = "".join(lines_old[i1:i2])
            txt_n = "".join(lines_new[j1:j2])
            if tag == 'equal':
                p.add_run(txt_o)
            elif tag == 'delete':
                apply_shading(p.add_run(txt_o), "F8D7DA")
            elif tag == 'replace':
                render_word_diff_inline(p, txt_o, txt_n)

        # Pure additions (not part of a replace)
        for tag, i1, i2, j1, j2 in refined_opcodes:
            if tag == 'insert':
                p.add_run(" [+] ")
                apply_shading(p.add_run("".join(lines_new[j1:j2])), "D4EDDA")
