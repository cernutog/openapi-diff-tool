from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime
import os
import difflib
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_COLOR_INDEX
from dependency_tracer import DependencyTracer

# --- OXML Helpers (Safe Insertion) ---
def get_or_add_child(parent, tag_name, order_list=None):
    if order_list is None:
        order_list = []
    child = parent.find(qn(tag_name))
    if child is not None:
        return child
    child = OxmlElement(tag_name)
    
    if order_list:
        try:
            my_idx = order_list.index(tag_name)
            for i in range(my_idx + 1, len(order_list)):
                next_tag = order_list[i]
                next_element = parent.find(qn(next_tag))
                if next_element is not None:
                    parent.insert(parent.index(next_element), child)
                    return child
        except ValueError:
            pass
            
    parent.append(child)
    return child

TBL_PR_ORDER = ['w:tblStyle', 'w:tblpPr', 'w:tblOverlap', 'w:bidiVisual', 'w:tblStyleRowBandSize', 'w:tblStyleColBandSize', 'w:tblW', 'w:jc', 'w:tblCellSpacing', 'w:tblInd', 'w:tblBorders', 'w:shd', 'w:tblLayout', 'w:tblCellMar', 'w:tblLook']
TC_PR_ORDER = ['w:cnfStyle', 'w:tcW', 'w:gridSpan', 'w:hMerge', 'w:vMerge', 'w:tcBorders', 'w:shd', 'w:noWrap', 'w:tcMar', 'w:textDirection', 'w:tcFitText', 'w:vAlign', 'w:hideMark', 'w:headers', 'w:cellIns', 'w:cellDel', 'w:cellMerge', 'w:tcPrChange']
P_PR_ORDER = ['w:pStyle', 'w:keepNext', 'w:keepLines', 'w:pageBreakBefore', 'w:framePr', 'w:widowControl', 'w:numPr', 'w:suppressLineNumbers', 'w:pBdr', 'w:shd', 'w:tabs', 'w:suppressAutoHyphens', 'w:kinsoku', 'w:wordWrap', 'w:overflowPunct', 'w:topLinePunct', 'w:autoSpaceDE', 'w:autoSpaceDN', 'w:bidi', 'w:adjustRightInd', 'w:snapToGrid', 'w:spacing', 'w:ind', 'w:contextualSpacing', 'w:mirrorIndents', 'w:suppressOverlap', 'w:jc', 'w:textDirection', 'w:textAlignment', 'w:textboxTightWrap', 'w:outlineLvl', 'w:divId', 'w:cnfStyle', 'w:rPr', 'w:sectPr', 'w:pPrChange']
R_PR_ORDER = ['w:rStyle', 'w:rFonts', 'w:b', 'w:bCs', 'w:i', 'w:iCs', 'w:caps', 'w:smallCaps', 'w:strike', 'w:dstrike', 'w:outline', 'w:shadow', 'w:emboss', 'w:imprint', 'w:noProof', 'w:snapToGrid', 'w:vanish', 'w:webHidden', 'w:color', 'w:spacing', 'w:w', 'w:kern', 'w:position', 'w:sz', 'w:szCs', 'w:highlight', 'w:u', 'w:effect', 'w:bdr', 'w:shd', 'w:fitText', 'w:vertAlign', 'w:rtl', 'w:cs', 'w:em', 'w:lang', 'w:eastAsianLayout', 'w:specVanish', 'w:oMath']

class AnalyticDocxGenerator:
    def __init__(self, spec1, spec2, diff, old_path=None, new_path=None, variables=None, template_path=None):
        self.spec1 = spec1
        self.spec2 = spec2
        self.diff = diff
        self.old_path = old_path
        self.new_path = new_path
        self.variables = variables or {}
        
        # Template Loading Logic
        # 1. Specific template passed in (e.g., template_analytic.docx)
        # 2. Fallback to generic 'template.docx'
        # 3. No template
        
        self.template_path = None
        self.has_template = False
        
        if template_path and os.path.exists(template_path):
            self.template_path = template_path
        elif os.path.exists("template.docx"):
            self.template_path = "template.docx"
            
        if self.template_path:
            self.doc = Document(self.template_path)
            self.has_template = True
        else:
            self.doc = Document()
            self.has_template = False
            
        # Initialize Dependency Tracer with NEW spec to find where schemas are NOW used
        self.tracer = DependencyTracer(spec2)
        self.tracer.resolve_transitive_impact()
            
        self._setup_styles()
        
        # Only apply default layout/header if NO template is provided
        if not self.has_template:
            self._setup_page_layout()
            self._add_header_footer()

    def _setup_styles(self):
        # Serif Title (Georgia)
        if 'Title' in self.doc.styles:
            style = self.doc.styles['Title']
            style.font.name = 'Georgia'
            style.font.size = Pt(26)
            style.font.bold = True
            style.font.color.rgb = RGBColor(31, 78, 121) # Dark Blue

        # Normal text
        normal = self.doc.styles['Normal']
        normal.font.name = 'Segoe UI'
        normal.font.size = Pt(10)
        normal.paragraph_format.space_after = Pt(8)
        normal.paragraph_format.line_spacing = 1.15
        normal.paragraph_format.left_indent = Pt(0) # Force reset indentation
        normal.paragraph_format.first_line_indent = Pt(0) # Force reset first line

        # Table Header (Clean - No Background)
        if 'Table Header' not in self.doc.styles:
            s = self.doc.styles.add_style('Table Header', WD_STYLE_TYPE.PARAGRAPH)
            s.base_style = normal
            s.font.bold = True
            s.font.color.rgb = RGBColor(80, 80, 80) # Dark Grey
            s.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            s.paragraph_format.space_before = Pt(6)
            s.paragraph_format.space_after = Pt(6)

        # Table Text
        if 'Table Text' not in self.doc.styles:
            s = self.doc.styles.add_style('Table Text', WD_STYLE_TYPE.PARAGRAPH)
            s.base_style = normal
            s.font.size = Pt(9)
            s.paragraph_format.space_before = Pt(4)
            s.paragraph_format.space_after = Pt(4)

        # Headings (Clean with bottom border)
        h1 = self.doc.styles['Heading 1']
        h1.font.name = 'Segoe UI'
        h1.font.size = Pt(16)
        h1.font.bold = True
        h1.font.color.rgb = RGBColor(0, 51, 102) # Dark Navy
        h1.paragraph_format.space_before = Pt(24)
        h1.paragraph_format.space_after = Pt(12)
        # Border for H1
        pPr = h1._element.get_or_add_pPr()
        pbdr = get_or_add_child(pPr, 'w:pBdr', P_PR_ORDER)
        pbdr.clear()
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '4')
        bottom.set(qn('w:color'), '003366')
        pbdr.append(bottom)

        h2 = self.doc.styles['Heading 2']
        h2.font.name = 'Segoe UI'
        h2.font.size = Pt(13)
        h2.font.bold = True
        h2.font.color.rgb = RGBColor(0, 51, 102)
        h2.paragraph_format.space_before = Pt(18)
        h2.paragraph_format.space_after = Pt(6)

        h3 = self.doc.styles['Heading 3']
        h3.font.name = 'Segoe UI'
        h3.font.size = Pt(11)
        h3.font.bold = True
        h3.font.color.rgb = RGBColor(51, 51, 51)

        # List Styles
        for i in range(2, 6):
            style_name = f'List Bullet {i}'
            if style_name not in self.doc.styles:
                s = self.doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
                s.base_style = self.doc.styles['List Bullet']
                s.font.name = 'Segoe UI'
                s.paragraph_format.left_indent = Inches(0.25 * i)
                s.paragraph_format.first_line_indent = Inches(-0.25)

    def _setup_page_layout(self):
        section = self.doc.sections[0]
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)

    def _add_spec_metadata(self):
        # Create a table for Spec Details
        widths = [Inches(1.5), Inches(2.75), Inches(2.75)]
        table = self._create_table(3, widths)
        
        # Custom Header Styling (Dark Blue Background, White Text)
        headers = ['Detail', 'Old Specification', 'New Specification']
        row = table.rows[0]
        for i, text in enumerate(headers):
            cell = row.cells[i]
            cell.text = text
            p = cell.paragraphs[0]
            p.style = 'Table Header'
            p.runs[0].font.color.rgb = RGBColor(255, 255, 255) # White
            
            # Dark Blue Background
            tcPr = cell._tc.get_or_add_tcPr()
            shd = get_or_add_child(tcPr, 'w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:fill'), '1F4E79') # Dark Blue
            
            # Borders
            tcBorders = get_or_add_child(tcPr, 'w:tcBorders')
            bottom = get_or_add_child(tcBorders, 'w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '12')
            bottom.set(qn('w:color'), 'FFFFFF')
        
        # Helper to get info
        def get_info(spec, path):
            info = spec.get('info', {})
            return {
                'file': os.path.basename(path) if path else "N/A",
                'title': info.get('title', 'N/A'),
                'version': info.get('version', 'N/A')
            }
            
        old_info = get_info(self.spec1, self.old_path)
        new_info = get_info(self.spec2, self.new_path)
        
        # Rows
        rows = [
            ("File Name", old_info['file'], new_info['file']),
            ("API Title", old_info['title'], new_info['title']),
            ("Version", old_info['version'], new_info['version'])
        ]
        
        for label, old_val, new_val in rows:
            row = table.add_row()
            row.cells[0].text = label
            row.cells[1].text = str(old_val)
            row.cells[2].text = str(new_val)
            
            # Style Label
            row.cells[0].paragraphs[0].runs[0].font.bold = True
            
            for cell in row.cells:
                self._style_body_cell(cell)
                
        self.doc.add_paragraph().paragraph_format.space_after = Pt(12)

    def _add_header_footer(self):
        section = self.doc.sections[0]
        # Header
        h_para = section.header.paragraphs[0]
        h_para.text = "OpenAPI Analytic Report"
        h_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        h_para.runs[0].font.size = Pt(8)
        h_para.runs[0].font.color.rgb = RGBColor(128, 128, 128)
        # Footer
        f_para = section.footer.paragraphs[0]
        f_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        f_para.text = f"Generated on {datetime.datetime.now().strftime('%Y-%m-%d')} | Page "
        f_para.runs[0].font.size = Pt(8)
        f_para.runs[0].font.color.rgb = RGBColor(128, 128, 128)
        self._add_page_number(f_para)

    def _add_page_number(self, paragraph):
        run = paragraph.add_run()
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        run._r.append(fldChar1)
        instrText = OxmlElement('w:instrText')
        instrText.text = "PAGE"
        run_instr = paragraph.add_run()
        run_instr._r.append(instrText)
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        run_end = paragraph.add_run()
        run_end._r.append(fldChar2)

    def _remove_all_borders(self, table):
        tbl = table._tbl
        tblPr = tbl.tblPr
        tblBorders = get_or_add_child(tblPr, 'w:tblBorders', TBL_PR_ORDER)
        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = get_or_add_child(tblBorders, f'w:{border_name}')
            border.set(qn('w:val'), 'nil')

    def _set_table_fixed_width(self, table, width_inches):
        tbl = table._tbl
        tblPr = tbl.tblPr
        tblW = get_or_add_child(tblPr, 'w:tblW', TBL_PR_ORDER)
        tblW.set(qn('w:w'), str(int(width_inches * 1440)))
        tblW.set(qn('w:type'), 'dxa')
        tblLayout = get_or_add_child(tblPr, 'w:tblLayout', TBL_PR_ORDER)
        tblLayout.set(qn('w:type'), 'fixed')

    def _create_table(self, cols, widths):
        table = self.doc.add_table(rows=1, cols=cols)
        self._remove_all_borders(table)
        
        # Calculate total width from columns (widths are Length objects)
        # sum(widths) returns int (EMUs), losing .inches attribute.
        # We need to sum the .inches values directly.
        total_width = sum(w.inches for w in widths)
        self._set_table_fixed_width(table, total_width)
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        
        # Set Grid
        tblPr = table._tblPr
        tblGrid = table._element.find(qn('w:tblGrid'))
        if tblGrid is None:
            tblGrid = OxmlElement('w:tblGrid')
            table._element.insert(table._element.index(tblPr) + 1, tblGrid)
        else:
            tblGrid.clear()
            
        for w in widths:
            col = OxmlElement('w:gridCol')
            col.set(qn('w:w'), str(int(w.twips)))
            tblGrid.append(col)
            
        for i, cell in enumerate(table.rows[0].cells):
            cell.width = widths[i]
            
        return table

    def _style_header_row(self, row, headers):
        for i, text in enumerate(headers):
            cell = row.cells[i]
            cell.text = text
            p = cell.paragraphs[0]
            p.style = 'Table Header'
            
            # Bottom border only
            tcPr = cell._tc.get_or_add_tcPr()
            tcBorders = get_or_add_child(tcPr, 'w:tcBorders', TC_PR_ORDER)
            bottom = get_or_add_child(tcBorders, 'w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '12')
            bottom.set(qn('w:color'), '000000')

    def _style_body_cell(self, cell):
        for p in cell.paragraphs:
            p.style = 'Table Text'
            
        # Horizontal Border Only (Light Grey)
        tcPr = cell._tc.get_or_add_tcPr()
        tcBorders = get_or_add_child(tcPr, 'w:tcBorders', TC_PR_ORDER)
        for side in ['top', 'left', 'right']:
            tag = get_or_add_child(tcBorders, f'w:{side}')
            tag.set(qn('w:val'), 'nil')
        bottom = get_or_add_child(tcBorders, 'w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '4')
        bottom.set(qn('w:color'), 'E0E0E0')

    def _add_pill_badge(self, paragraph, text, color_override=None):
        # Add spacing
        run = paragraph.add_run(f"  {text}  ")
        run.font.size = Pt(7)
        run.font.bold = True
        run.font.name = 'Segoe UI'
        
        rPr = run._r.get_or_add_rPr()
        shd = get_or_add_child(rPr, 'w:shd')
        shd.set(qn('w:val'), 'clear')
        
        # Logic for colors
        bg_color = 'E2E3E5' # Default Grey
        text_color = RGBColor(56, 61, 65)
        
        if color_override:
             # If caller provided a specific color (e.g. from old generator), map it to pastel
             # Old: Green (28A745) -> Pastel Green
             if color_override == '28A745' or color_override == '28a745': # Green
                 bg_color = 'D4EDDA'
                 text_color = RGBColor(21, 87, 36)
             elif color_override == 'DC3545' or color_override == 'dc3545': # Red
                 bg_color = 'F8D7DA'
                 text_color = RGBColor(114, 28, 36)
             elif color_override == '17A2B8' or color_override == '17a2b8': # Cyan
                 bg_color = 'D1ECF1'
                 text_color = RGBColor(12, 84, 96)
             elif color_override == 'FFC107' or color_override == 'ffc107': # Yellow
                 bg_color = 'FFF3CD'
                 text_color = RGBColor(133, 100, 4)
             elif color_override == 'NEUTRAL': # Neutral Grey
                 bg_color = 'E0E0E0'
                 text_color = RGBColor(80, 80, 80)
        else:
            # Text-based mapping
            if 'NEW' in text or 'ADDED' in text:
                 bg_color = 'D4EDDA'
                 text_color = RGBColor(21, 87, 36)
            elif 'REMOVED' in text or 'DELETED' in text:
                 bg_color = 'F8D7DA'
                 text_color = RGBColor(114, 28, 36)
            elif 'MODIFIED' in text or 'CHANGED' in text:
                 bg_color = 'FFF3CD'
                 text_color = RGBColor(133, 100, 4)
        
        shd.set(qn('w:fill'), bg_color)
        run.font.color.rgb = text_color
        
        # Add a small margin run to make the badge look wider
        paragraph.add_run(" ")

    def _add_legend(self):
        self.doc.add_heading('Legend of Changes', 1)
        
        table = self.doc.add_table(rows=0, cols=2)
        self._remove_all_borders(table)
        self._set_table_fixed_width(table, 7.0)
        
        legend_items = [
            ("ADDED", "28A745", "New component found in the new specification."),
            ("REMOVED", "DC3545", "Component from the old specification that is no longer present."),
            ("MODIFIED", "FFC107", "Component that exists in both versions but has changed content."),
            ("RENAMED", "17A2B8", "Component that has changed name but is structurally identical to a removed component."),
            ("RENAMED & MODIFIED", "FFC107", "Component that has been renamed and also has content changes."),
            ("REF RENAMED", "17A2B8", "A reference that points to a renamed component.")
        ]


        for label, color, desc in legend_items:
            row = table.add_row()
            # Badge Cell
            cell_badge = row.cells[0]
            cell_badge.width = Inches(1.5)
            cell_badge.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p_badge = cell_badge.paragraphs[0]
            p_badge.paragraph_format.space_before = Pt(0)
            p_badge.paragraph_format.space_after = Pt(0)
            self._add_pill_badge(p_badge, label, color)
            
            # Description Cell
            cell_desc = row.cells[1]
            cell_desc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p_desc = cell_desc.paragraphs[0]
            p_desc.paragraph_format.space_before = Pt(0)
            p_desc.paragraph_format.space_after = Pt(0)
            p_desc.text = desc
            p_desc.style = 'Table Text'
            
            # Apply borders
            for cell in row.cells:
                self._style_body_cell(cell)

    def generate(self, output_path):
        # Title
        self.doc.add_heading('OpenAPI Comparison - Analytical Report', 0)
        
        self._add_spec_metadata()
        


        self._add_legend()

        self._add_dashboard()
        self._add_general_info()
        self._add_endpoints()
        self._add_components()
        
        # Variable Substitution (Final Step)
        self._process_template_variables()
        
        self.doc.save(output_path)
        print(f"Analytic Report generated at {output_path}")

    def _process_template_variables(self):
        """
        Replaces {{ variable }} placeholders in the document with values.
        Prioritizes:
        1. Dynamic Variables (date, time, filenames)
        2. User Static Variables (from Preferences)
        """
        import datetime
        import getpass
        import sys
        
        # 1. Prepare Variables
        context = self.variables.copy()
        
        # Dynamic Defaults
        now = datetime.datetime.now()
        context['date'] = now.strftime('%Y-%m-%d')
        context['time'] = now.strftime('%H:%M')
        context['datetime'] = now.strftime('%Y-%m-%d %H:%M:%S')
        context['original_spec'] = os.path.basename(self.old_path) if self.old_path else "N/A"
        context['new_spec'] = os.path.basename(self.new_path) if self.new_path else "N/A"
        
        # Enriched Variables
        try:
            context['user'] = getpass.getuser()
        except:
            context['user'] = "Unknown"
            
        context['platform'] = sys.platform
        context['tool_version'] = "1.0.0"
        
        def get_size(path):
            if path and os.path.exists(path):
                size_bytes = os.path.getsize(path)
                if size_bytes < 1024: return f"{size_bytes} B"
                elif size_bytes < 1024*1024: return f"{size_bytes/1024:.1f} KB"
                else: return f"{size_bytes/(1024*1024):.1f} MB"
            return "N/A"
            
        context['file_size_old'] = get_size(self.old_path)
        context['file_size_new'] = get_size(self.new_path)
        
        # 2. Iterate and Replace
        # Body Paragraphs
        for p in self.doc.paragraphs:
            self._replace_text_in_paragraph(p, context)
            
        # Tables
        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        self._replace_text_in_paragraph(p, context)
                        
        # Headers & Footers
        for section in self.doc.sections:
            # Header
            for p in section.header.paragraphs:
                self._replace_text_in_paragraph(p, context)
            for table in section.header.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            self._replace_text_in_paragraph(p, context)
                            
            # Footer
            for p in section.footer.paragraphs:
                self._replace_text_in_paragraph(p, context)
            for table in section.footer.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            self._replace_text_in_paragraph(p, context)

    def _replace_text_in_paragraph(self, paragraph, context):
        if '{{' not in paragraph.text:
            return
            
        # Naive replacement (works for simple cases where {{var}} is in one run)
        # For robust replacement across runs, we'd need a more complex parser.
        # Assuming user inputs {{ var }} cleanly.
        
        text = paragraph.text
        for key, value in context.items():
            placeholder = f"{{{{ {key} }}}}"
            placeholder_tight = f"{{{{{key}}}}}" # Handle {{key}} without spaces
            
            if placeholder in text or placeholder_tight in text:
                text = text.replace(placeholder, str(value))
                text = text.replace(placeholder_tight, str(value))
        
        # If text changed, update runs. 
        # WARNING: This destroys formatting if runs were split.
        # Better approach: Iterate runs and replace if full match found.
        # But split runs (e.g. {{ var }}) are hard.
        # Simple approach: If match found, clear p and add new run with text.
        # Preserving style is hard.
        
        # Improved approach: Check each run.
        for run in paragraph.runs:
            for key, value in context.items():
                placeholder = f"{{{{ {key} }}}}"
                placeholder_tight = f"{{{{{key}}}}}"
                
                if placeholder in run.text:
                    run.text = run.text.replace(placeholder, str(value))
                if placeholder_tight in run.text:
                    run.text = run.text.replace(placeholder_tight, str(value))

    def _add_dashboard(self):
        self.doc.add_heading('Change Matrix', 1)
        
        # Single Compact Matrix Table
        # Columns: Category | New | Removed | Modified | Renamed
        widths = [Inches(2.0), Inches(1.25), Inches(1.25), Inches(1.25), Inches(1.25)]
        table = self._create_table(5, widths)
        
        # Header
        self._style_header_row(table.rows[0], ['Category', 'New', 'Removed', 'Modified', 'Renamed'])
        
        # Data Rows
        # Logical Order: Core -> Data -> I/O -> Auth -> Meta -> Advanced
        # Calculate Schema Metrics (Strict Partition)
        s_new = len(self.diff.new_components.get('schemas', []))
        s_rem = len(self.diff.removed_components.get('schemas', []))
        
        # Partition Modified vs Renamed
        # Renamed = In renamed_components AND (NOT in modified OR !is_substantial)
        # Modified = In modified_components AND is_substantial
        
        s_mod_count = 0
        s_ren_count = 0
        
        mod_items = self.diff.modified_components.get('schemas', {})
        ren_items = self.diff.renamed_components.get('schemas', {})
        
        # 1. Count Modified (Substantial)
        for s_name, changes in mod_items.items():
            if self._is_substantial_modification(s_name, changes, ren_items):
                s_mod_count += 1
        
        # 2. Count Renamed (Pure + Ref-Only)
        for old_name in ren_items:
            # If it's NOT in modified, it's pure rename
            if old_name not in mod_items:
                s_ren_count += 1
            else:
                # If it IS in modified, check if substantial
                if not self._is_substantial_modification(old_name, mod_items[old_name], ren_items):
                    s_ren_count += 1

        metrics = [
            ("Endpoints", len(self.diff.new_paths), len(self.diff.removed_paths), len(self.diff.modified_paths), 0),
            ("Schemas", s_new, s_rem, s_mod_count, s_ren_count),
            ("Parameters", len(self.diff.new_components.get('parameters', [])), len(self.diff.removed_components.get('parameters', [])), len(self.diff.modified_components.get('parameters', {})), 0),
            ("Responses", len(self.diff.new_components.get('responses', [])), len(self.diff.removed_components.get('responses', [])), len(self.diff.modified_components.get('responses', {})), 0),
            ("Headers", len(self.diff.new_components.get('headers', [])), len(self.diff.removed_components.get('headers', [])), len(self.diff.modified_components.get('headers', {})), 0),
            ("Security Schemes", len(self.diff.new_components.get('securitySchemes', [])), len(self.diff.removed_components.get('securitySchemes', [])), len(self.diff.modified_components.get('securitySchemes', {})), 0),
            ("Tags", len(self.diff.tags_changes.get('new', [])), len(self.diff.tags_changes.get('removed', [])), len(self.diff.tags_changes.get('modified', {})), 0),
            ("Servers", len(self.diff.servers_changes.get('new', [])), len(self.diff.servers_changes.get('removed', [])), len(self.diff.servers_changes.get('modified', {})), 0),
            ("Links", len(self.diff.new_components.get('links', [])), len(self.diff.removed_components.get('links', [])), len(self.diff.modified_components.get('links', {})), 0),
            ("Callbacks", len(self.diff.new_components.get('callbacks', [])), len(self.diff.removed_components.get('callbacks', [])), len(self.diff.modified_components.get('callbacks', {})), 0),
            ("Examples", len(self.diff.new_components.get('examples', [])), len(self.diff.removed_components.get('examples', [])), len(self.diff.modified_components.get('examples', {})), 0)
        ]
        
        for name, new_c, rem_c, mod_c, ren_c in metrics:
            # Skip row if all counts are zero
            if (new_c + rem_c + mod_c + ren_c) == 0:
                continue
                
            row = table.add_row()
            row.cells[0].text = name
            row.cells[1].text = str(new_c)
            row.cells[2].text = str(rem_c)
            row.cells[3].text = str(mod_c)
            row.cells[4].text = str(ren_c)
            
            # Style Category Cell
            p_cat = row.cells[0].paragraphs[0]
            p_cat.style = 'Table Text'
            p_cat.runs[0].font.bold = True
            p_cat.runs[0].font.color.rgb = RGBColor(50, 50, 50)
            
            # Style Number Cells
            for i in range(1, 5):
                cell = row.cells[i]
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.style = 'Table Text'
                val = int(cell.text)
                if val > 0:
                    p.runs[0].font.bold = True
                    # Optional: Color code?
                    if i == 1: # New
                        p.runs[0].font.color.rgb = RGBColor(40, 167, 69) # Green
                    elif i == 2: # Removed
                        p.runs[0].font.color.rgb = RGBColor(220, 53, 69) # Red
                    elif i == 3: # Modified
                        p.runs[0].font.color.rgb = RGBColor(180, 130, 0) # Amber
                    elif i == 4: # Renamed
                        p.runs[0].font.color.rgb = RGBColor(23, 162, 184) # Cyan
                else:
                    p.runs[0].font.color.rgb = RGBColor(200, 200, 200) # Light Grey for zeros
            
            # Apply borders
            for cell in row.cells:
                self._style_body_cell(cell)
                
        self.doc.add_paragraph().paragraph_format.space_after = Pt(24)

    def _add_general_info(self):
        self.doc.add_heading('General Info', 1)
        if not self.diff.info_changes:
            self.doc.add_paragraph('No changes in General Info.')
            return

        # Total avail = 6.4 (no indent here)
        total_avail = 6.4
        col1 = 1.2
        col23 = (total_avail - col1) / 2
        widths = [Inches(col1), Inches(col23), Inches(col23)]
        table = self._create_table(3, widths)
        self._style_header_row(table.rows[0], ['Field', 'Old Value', 'New Value'])
        
        for key, change in self.diff.info_changes.items():
            row = table.add_row()
            row.cells[0].text = str(key)
            
            if str(key).lower() == 'description':
                self._render_rich_diff(row.cells[1].paragraphs[0], row.cells[2].paragraphs[0], change['old'], change['new'])
            else:
                row.cells[1].text = str(change['old'])
                row.cells[2].text = str(change['new'])
            
            for cell in row.cells:
                self._style_body_cell(cell)

    def _add_endpoints(self):
        self.doc.add_heading('Endpoints Summary', 1)
        
        # New
        self.doc.add_heading('New Endpoints', 2)
        if self.diff.new_paths:
            # Sort based on new spec order
            new_order = list(self.spec2.get('paths', {}).keys())
            ordered_new = sorted(self.diff.new_paths, 
                                 key=lambda x: new_order.index(x) if x in new_order else 9999)
            for path in ordered_new:
                p = self.doc.add_paragraph(style='List Bullet')
                self._add_pill_badge(p, "NEW", "28A745")
                p.add_run(f" Endpoint: {path}")
        else:
            self.doc.add_paragraph('No new endpoints.')

        # Removed
        self.doc.add_heading('Removed Endpoints', 2)
        if self.diff.removed_paths:
            # Sort based on old spec order
            old_order = list(self.spec1.get('paths', {}).keys())
            ordered_removed = sorted(self.diff.removed_paths, 
                                     key=lambda x: old_order.index(x) if x in old_order else 9999)
            for path in ordered_removed:
                p = self.doc.add_paragraph(style='List Bullet')
                self._add_pill_badge(p, "REMOVED", "DC3545")
                p.add_run(f" Endpoint: {path}")
        else:
            self.doc.add_paragraph('No removed endpoints.')

        # Modified
        self.doc.add_heading('Modified Endpoints', 2)
        if self.diff.modified_paths:
            # Sort modified paths based on their order in the old specification
            original_order = list(self.spec1.get('paths', {}).keys())
            ordered_modified = sorted(self.diff.modified_paths.keys(), 
                                      key=lambda x: original_order.index(x) if x in original_order else 9999)
            
            for path in ordered_modified:
                changes = self.diff.modified_paths[path]
                self.doc.add_heading(path, 3)
                
                if 'new_ops' in changes:
                    p = self.doc.add_paragraph()
                    p.paragraph_format.left_indent = Inches(0.25)
                    self._add_pill_badge(p, "NEW", "28A745")
                    p.add_run(" Method: " + f"{', '.join(changes['new_ops']).upper()}")
                
                if 'removed_ops' in changes:
                    p = self.doc.add_paragraph()
                    p.paragraph_format.left_indent = Inches(0.25)
                    self._add_pill_badge(p, "REMOVED", "DC3545")
                    p.add_run(" Method: " + f"{', '.join(changes['removed_ops']).upper()}")
                
                if 'modified_ops' in changes:
                    p = self.doc.add_paragraph()
                    p.paragraph_format.left_indent = Inches(0.25)
                    self._add_pill_badge(p, "MODIFIED", "FFC107")
                    p.add_run(" Method: The following methods have changed:")
                    
                    for op, op_changes in changes['modified_ops'].items():
                        p_op = self.doc.add_paragraph()
                        p_op.paragraph_format.left_indent = Inches(0.5)
                        self._add_pill_badge(p_op, op.upper(), "6C757D")
                        
                        # Check for Metadata Changes (summary, description, etc.)
                        meta_changes = {k: v for k, v in op_changes.items() if k in ['summary', 'description', 'deprecated', 'operationId']}
                        if meta_changes:
                            p_meta = self.doc.add_paragraph('Operation Metadata Changes:')
                            p_meta.paragraph_format.left_indent = Inches(0.75)
                            
                            # Adjusted widths (Total 6.4 - 0.75 indent = 5.65)
                            total_avail = 6.4 - 0.75
                            col1 = 1.2
                            col23 = (total_avail - col1) / 2
                            widths = [Inches(col1), Inches(col23), Inches(col23)]
                            table = self._create_table(3, widths)
                            tblPr = table._tblPr
                            tblInd = get_or_add_child(tblPr, 'w:tblInd', TBL_PR_ORDER)
                            tblInd.set(qn('w:w'), str(int(Inches(0.75).twips)))
                            tblInd.set(qn('w:type'), 'dxa')
                            
                            self._style_header_row(table.rows[0], ['Field', 'Old Value', 'New Value'])
                            
                            for k, v in meta_changes.items():
                                row = table.add_row()
                                row.cells[0].text = k
                                if k == 'description':
                                    self._render_rich_diff(row.cells[1].paragraphs[0], row.cells[2].paragraphs[0], v['old'], v['new'])
                                else:
                                    row.cells[1].text = str(v['old'] or '')
                                    row.cells[2].text = str(v['new'] or '')
                                for cell in row.cells:
                                    self._style_body_cell(cell)
                            self.doc.add_paragraph().paragraph_format.space_after = Pt(6)

                        if 'parameters' in op_changes:
                            self._add_parameter_changes(op_changes['parameters'])
                        if 'requestBody' in op_changes:
                            self._add_request_body_changes(op_changes['requestBody'])
                        if 'responses' in op_changes:
                            self._add_response_changes(op_changes['responses'])
        else:
            self.doc.add_paragraph('No modified endpoints.')

    def _add_parameter_changes(self, params_diff):
        if 'new' in params_diff:
            for param in params_diff['new']:
                p = self.doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.75)
                self._add_pill_badge(p, "NEW", "28A745")
                p.add_run(" Param: " + param)
        if 'removed' in params_diff:
            for param in params_diff['removed']:
                p = self.doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.75)
                self._add_pill_badge(p, "REMOVED", "DC3545")
                p.add_run(" Param: " + param)
        if 'modified' in params_diff:
            for param, changes in params_diff['modified'].items():
                p = self.doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.75)
                self._add_pill_badge(p, "MODIFIED", "FFC107")
                p.add_run(f" Param: {param}")
                
                # Split changes into attributes and schema
                attr_changes = {k: v for k, v in changes.items() if k != 'schema' and isinstance(v, dict) and 'old' in v}
                
                if attr_changes:
                    # Render attribute changes in a table (Total 6.4 - 0.75 indent = 5.65)
                    total_avail = 6.4 - 0.75
                    col1 = 1.2
                    col23 = (total_avail - col1) / 2
                    widths = [Inches(col1), Inches(col23), Inches(col23)]
                    table = self._create_table(3, widths)
                    tblPr = table._tblPr
                    tblInd = get_or_add_child(tblPr, 'w:tblInd', TBL_PR_ORDER)
                    tblInd.set(qn('w:w'), str(int(Inches(0.75).twips)))
                    tblInd.set(qn('w:type'), 'dxa')
                    
                    self._style_header_row(table.rows[0], ['Field', 'Old Value', 'New Value'])
                    
                    for k, v in attr_changes.items():
                        row = table.add_row()
                        row.cells[0].text = k
                        if k == 'description':
                            self._render_rich_diff(row.cells[1].paragraphs[0], row.cells[2].paragraphs[0], v['old'], v['new'])
                        else:
                            row.cells[1].text = str(v['old'] or '')
                            row.cells[2].text = str(v['new'] or '')
                        for cell in row.cells:
                            self._style_body_cell(cell)
                    self.doc.add_paragraph().paragraph_format.space_after = Pt(6)

                if 'schema' in changes:
                    self._render_schema_diff_details(changes['schema'], indent_level=1.0)

    def _add_request_body_changes(self, rb_diff):
        p = self.doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.75)
        # Removed the "DETAILS" pill badge
        run = p.add_run("Request Body:")
        run.font.bold = True

        if 'required' in rb_diff:
            val = rb_diff['required']
            p_req = self.doc.add_paragraph()
            p_req.paragraph_format.left_indent = Inches(1.0)
            p_req.add_run(f"Required changed from {val['old']} to {val['new']}")

        if 'content' in rb_diff:
            content_diff = rb_diff['content']
            if 'modified' in content_diff:
                for media_type, mt_changes in content_diff['modified'].items():
                    p_mt = self.doc.add_paragraph()
                    p_mt.paragraph_format.left_indent = Inches(1.0)
                    p_mt.add_run(f"Media Type: {media_type}")
                    
                    if 'schema' in mt_changes:
                        p_note = self.doc.add_paragraph('Schema Details (Old vs New):')
                        p_note.paragraph_format.left_indent = Inches(1.1)
                        p_note.runs[0].font.italic = True
                        self._render_schema_diff_details(mt_changes['schema'], indent_level=1.25)
                    
                    if 'examples' in mt_changes:
                         self._add_examples_changes_section(mt_changes['examples'], indent_level=1.1)

                    # Extensions and other attributes
                    mt_attr = {k: v for k, v in mt_changes.items() if k not in ['schema', 'examples'] and (k.startswith('x-') or (isinstance(v, dict) and 'old' in v))}
                    if mt_attr:
                         self._add_metadata_table(mt_attr, indent_level=1.1)

        # Body-level extensions
        rb_ext = {k: v for k, v in rb_diff.items() if k not in ['required', 'content'] and k.startswith('x-')}
        if rb_ext:
             self._add_metadata_table(rb_ext, indent_level=1.0)

    def _add_response_changes(self, responses_diff):
        if 'new' in responses_diff:
            for code in responses_diff['new']:
                p = self.doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.75)
                self._add_pill_badge(p, "NEW", "28A745")
                p.add_run(f" Response {code}")
        if 'removed' in responses_diff:
            for code in responses_diff['removed']:
                p = self.doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.75)
                self._add_pill_badge(p, "REMOVED", "DC3545")
                p.add_run(f" Response {code}")
        if 'modified' in responses_diff:
            for code, changes in responses_diff['modified'].items():
                p = self.doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.75)
                self._add_pill_badge(p, "MODIFIED", "FFC107")
                p.add_run(f" Response {code}")
                
                # Split changes into attributes and others (content, headers)
                # Now including x- extensions in attr_changes
                attr_changes = {k: v for k, v in changes.items() if k not in ['content', 'headers', 'examples'] and (k.startswith('x-') or (isinstance(v, dict) and 'old' in v))}
                
                if attr_changes:
                     # Calculate available width (Total 6.4 - 0.75 indent = 5.65)
                     total_avail = 6.4 - 0.75
                     col1 = 1.2
                     col23 = (total_avail - col1) / 2
                     widths = [Inches(col1), Inches(col23), Inches(col23)]
                     table = self._create_table(3, widths)
                     tblPr = table._tblPr
                     tblInd = get_or_add_child(tblPr, 'w:tblInd', TBL_PR_ORDER)
                     tblInd.set(qn('w:w'), str(int(Inches(0.75).twips)))
                     tblInd.set(qn('w:type'), 'dxa')
                     
                     self._style_header_row(table.rows[0], ['Field', 'Old Value', 'New Value'])
                     
                     for k, v in attr_changes.items():
                         row = table.add_row()
                         row.cells[0].text = k
                         if k == 'description':
                             self._render_rich_diff(row.cells[1].paragraphs[0], row.cells[2].paragraphs[0], v.get('old'), v.get('new'))
                         else:
                             row.cells[1].text = str(v.get('old') if isinstance(v, dict) else v)
                             row.cells[2].text = str(v.get('new') if isinstance(v, dict) else v)
                         for cell in row.cells:
                             self._style_body_cell(cell)
                     self.doc.add_paragraph().paragraph_format.space_after = Pt(6)

                if 'content' in changes:
                    content_diff = changes['content']
                    # Added Media Types
                    if 'new' in content_diff:
                        for mt in content_diff['new']:
                            p_mt = self.doc.add_paragraph()
                            p_mt.paragraph_format.left_indent = Inches(1.0)
                            self._add_pill_badge(p_mt, "NEW", "28A745")
                            p_mt.add_run(f" Media Type: {mt}")
                    
                    # Removed Media Types
                    if 'removed' in content_diff:
                        for mt in content_diff['removed']:
                            p_mt = self.doc.add_paragraph()
                            p_mt.paragraph_format.left_indent = Inches(1.0)
                            self._add_pill_badge(p_mt, "REMOVED", "DC3545")
                            p_mt.add_run(f" Media Type: {mt}")

                    if 'modified' in content_diff:
                        for media_type, mt_changes in content_diff['modified'].items():
                            p_mt = self.doc.add_paragraph()
                            p_mt.paragraph_format.left_indent = Inches(1.0)
                            p_mt.add_run(f"Media Type: {media_type}")
                            
                            if 'schema' in mt_changes:
                                # Clarity note if it's a schema change
                                p_note = self.doc.add_paragraph('Schema Details (Old vs New):')
                                p_note.paragraph_format.left_indent = Inches(1.1)
                                p_note.runs[0].font.italic = True
                                self._render_schema_diff_details(mt_changes['schema'], indent_level=1.25)
                            
                            if 'examples' in mt_changes:
                                # Dedicated example change section
                                self._add_examples_changes_section(mt_changes['examples'], indent_level=1.1)

                            # Other media type attributes (extensions, encoding, etc.)
                            mt_attr = {k: v for k, v in mt_changes.items() if k not in ['schema', 'examples'] and (k.startswith('x-') or (isinstance(v, dict) and 'old' in v))}
                            if mt_attr:
                                self._add_metadata_table(mt_attr, indent_level=1.1, title="Media Type Metadata Changes:")
                                
                if 'headers' in changes:
                    h_diff = changes['headers']
                    # Added Headers
                    if 'new' in h_diff:
                        for h_name in h_diff['new']:
                            p_hdr = self.doc.add_paragraph()
                            p_hdr.paragraph_format.left_indent = Inches(1.0)
                            self._add_pill_badge(p_hdr, "NEW", "28A745")
                            p_hdr.add_run(f" Header: {h_name}")
                    
                    # Removed Headers
                    if 'removed' in h_diff:
                        for h_name in h_diff['removed']:
                            p_hdr = self.doc.add_paragraph()
                            p_hdr.paragraph_format.left_indent = Inches(1.0)
                            self._add_pill_badge(p_hdr, "REMOVED", "DC3545")
                            p_hdr.add_run(f" Header: {h_name}")

                    # Modified headers
                    if 'modified' in h_diff:
                        for h_name, h_changes in h_diff['modified'].items():
                            p_hdr = self.doc.add_paragraph()
                            p_hdr.paragraph_format.left_indent = Inches(1.0)
                            self._add_pill_badge(p_hdr, "MODIFIED", "FFC107")
                            p_hdr.add_run(f" Header: {h_name}")
                            # Attributes (including x- extensions)
                            h_attr = {k: v for k, v in h_changes.items() if k != 'schema' and (k.startswith('x-') or (isinstance(v, dict) and 'old' in v))}
                            if h_attr:
                                self._add_metadata_table(h_attr, indent_level=1.25)
                            if 'schema' in h_changes:
                                self._render_schema_diff_details(h_changes['schema'], indent_level=1.25)

    def _add_metadata_table(self, attr_changes, indent_level=0.5, title=None):
        if not attr_changes: return
        if title:
            p = self.doc.add_paragraph(title)
            p.paragraph_format.left_indent = Inches(indent_level)
            p.runs[0].font.bold = True

        # Calculate available width (Max 6.5" total - indent)
        total_avail = 6.5 - indent_level
        col1 = 1.2 # Reduced from 1.5 to save space
        col23 = (total_avail - col1) / 2
        widths = [Inches(col1), Inches(col23), Inches(col23)]
        
        table = self._create_table(3, widths)
        tblInd = get_or_add_child(table._tblPr, 'w:tblInd', TBL_PR_ORDER)
        tblInd.set(qn('w:w'), str(int(Inches(indent_level).twips)))
        tblInd.set(qn('w:type'), 'dxa')
        self._style_header_row(table.rows[0], ['Field', 'Old Value', 'New Value'])
        for k, v in attr_changes.items():
            row = table.add_row()
            row.cells[0].text = k
            if k == 'description':
                self._render_rich_diff(row.cells[1].paragraphs[0], row.cells[2].paragraphs[0], v.get('old'), v.get('new'))
            else:
                row.cells[1].text = str(v.get('old') if isinstance(v, dict) else v)
                row.cells[2].text = str(v.get('new') if isinstance(v, dict) else v)
            for c in row.cells: self._style_body_cell(c)
        self.doc.add_paragraph().paragraph_format.space_after = Pt(6)

    def _add_examples_changes_section(self, exs_diff, indent_level=1.0):
        # Added
        if 'new' in exs_diff:
            for ex in exs_diff['new']:
                p = self.doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(indent_level)
                self._add_pill_badge(p, "NEW", "28A745")
                p.add_run(f" Example: {ex}")
        # Removed
        if 'removed' in exs_diff:
            for ex in exs_diff['removed']:
                p = self.doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(indent_level)
                self._add_pill_badge(p, "REMOVED", "DC3545")
                p.add_run(f" Example: {ex}")
        # Modified
        if 'modified' in exs_diff:
            for ex, changes in exs_diff['modified'].items():
                p = self.doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(indent_level)
                self._add_pill_badge(p, "MODIFIED", "FFC107")
                p.add_run(f" Example: {ex}")
                self._add_metadata_table(changes, indent_level=indent_level + 0.25)

    def _render_schema_diff_details(self, changes, indent_level=0.5):
        # 1. $ref Change
        if '$ref' in changes:
            val = changes['$ref']
            p = self.doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(indent_level)
            self._add_pill_badge(p, "MODIFIED", "FFC107")
            p.add_run(f" Schema Ref changed from '{val['old']}' to '{val['new']}'")
            return

        # 2. Schema Constraint Changes (Metadata)
        # Exclude structural things handled elsewhere
        structural = ['properties', 'items', 'allOf', 'oneOf', 'anyOf', 'not', 'additionalProperties', '$ref']
        attr_changes = {k: v for k, v in changes.items() if k not in structural and isinstance(v, dict) and 'old' in v}
        
        if attr_changes:
            # Calculate available width (Max 6.5" total - indent)
            total_avail = 6.5 - indent_level
            col1 = 1.4 # Reduced from 1.5
            col23 = (total_avail - col1) / 2
            widths = [Inches(col1), Inches(col23), Inches(col23)]
            
            table = self._create_table(3, widths)
            tblInd = get_or_add_child(table._tblPr, 'w:tblInd', TBL_PR_ORDER)
            tblInd.set(qn('w:w'), str(int(Inches(indent_level).twips)))
            tblInd.set(qn('w:type'), 'dxa')
            
            self._style_header_row(table.rows[0], ['Schema Attribute', 'Old Value', 'New Value'])
            for k, v in attr_changes.items():
                row = table.add_row()
                row.cells[0].text = k
                if k == 'description':
                    self._render_rich_diff(row.cells[1].paragraphs[0], row.cells[2].paragraphs[0], v['old'], v['new'])
                else:
                    row.cells[1].text = str(v['old'] or '')
                    row.cells[2].text = str(v['new'] or '')
                for cell in row.cells: self._style_body_cell(cell)
            self.doc.add_paragraph().paragraph_format.space_after = Pt(6)

        # 3. Properties Table
        if 'properties' in changes and 'modified' in changes['properties']:
            p_prop = self.doc.add_paragraph('Property-Level Changes:')
            p_prop.paragraph_format.left_indent = Inches(indent_level)
            
            # Calculate available width (Max 6.5" total - indent)
            total_avail = 6.5 - indent_level
            col1 = 1.3 # Reduced from 1.5
            col2 = 0.7 # Reduced from 0.8
            col34 = (total_avail - col1 - col2) / 2
            
            widths = [Inches(col1), Inches(col2), Inches(col34), Inches(col34)]
            
            # Indent table
            table = self._create_table(4, widths)
            tblPr = table._tblPr
            tblInd = get_or_add_child(tblPr, 'w:tblInd', TBL_PR_ORDER)
            tblInd.set(qn('w:w'), str(int(Inches(indent_level).twips)))
            tblInd.set(qn('w:type'), 'dxa')
            
            self._style_header_row(table.rows[0], ['Property', 'Change', 'Old', 'New'])
            
            for prop, p_diff in changes['properties']['modified'].items():
                row = table.add_row()
                row.cells[0].text = prop
                cell_old = row.cells[2]
                cell_new = row.cells[3]
                p_old = cell_old.paragraphs[0]
                p_new = cell_new.paragraphs[0]
                
                for k, v in p_diff.items():
                    if k == 'description' and isinstance(v, dict) and 'old' in v:
                        p_old.add_run("description: ")
                        p_new.add_run("description: ")
                        self._render_rich_diff(p_old, p_new, v['old'], v['new'])
                        p_old.add_run("\n")
                        p_new.add_run("\n")
                    elif isinstance(v, dict) and 'old' in v:
                        old_val = v['old']
                        new_val = v['new']
                        if isinstance(old_val, (dict, list)):
                            import json
                            old_val = json.dumps(old_val, indent=2)
                        if isinstance(new_val, (dict, list)):
                            import json
                            new_val = json.dumps(new_val, indent=2)
                        p_old.add_run(f"{k}: {old_val}\n")
                        p_new.add_run(f"{k}: {new_val}\n")
                    else:
                        import json
                        p_old.add_run(f"{k}: {json.dumps(v, default=str)}\n")
                        p_new.add_run(f"{k}: {json.dumps(v, default=str)}\n")
                
                for cell in row.cells:
                    self._style_body_cell(cell)
            
            self.doc.add_paragraph().paragraph_format.space_after = Pt(6)

        # 3. Property Additions/Removals
        if 'properties' in changes:
            if 'new' in changes['properties'] and changes['properties']['new']:
                for prop in changes['properties']['new']:
                    p = self.doc.add_paragraph()
                    p.paragraph_format.left_indent = Inches(indent_level)
                    self._add_pill_badge(p, "NEW", "28A745")
                    p.add_run(f" Property: {prop}")
                
            if 'removed' in changes['properties'] and changes['properties']['removed']:
                for prop in changes['properties']['removed']:
                    p = self.doc.add_paragraph()
                    p.paragraph_format.left_indent = Inches(indent_level)
                    self._add_pill_badge(p, "REMOVED", "DC3545")
                    p.add_run(f" Property: {prop}")

        # 4. Other Schema Changes
        for key, val in changes.items():
            if key in ['properties', '$ref']:
                continue
                
            p = self.doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(indent_level)
            
            # Action badge
            action = "MODIFIED"
            color = "FFC107"
            if isinstance(val, dict) and 'added' in val and not val.get('removed'): 
                action, color = "NEW", "28A745"
            elif isinstance(val, dict) and 'removed' in val and not val.get('added'):
                action, color = "REMOVED", "DC3545"
                
            self._add_pill_badge(p, action, color)
            p.add_run(f" {key.upper()}: ")
            
            if isinstance(val, dict) and ('added' in val or 'removed' in val):
                 if 'added' in val and val['added']:
                    p.add_run("Added options:")
                    for item in val['added']:
                        p_sub = self.doc.add_paragraph()
                        p_sub.paragraph_format.left_indent = Inches(indent_level + 0.25)
                        self._add_pill_badge(p_sub, "ADDED", "28A745")
                        if isinstance(item, dict):
                             p_sub.add_run(self._format_schema_summary(item))
                        else:
                             p_sub.add_run(str(item))
                 if 'removed' in val and val['removed']:
                    if 'added' in val and val['added']:
                        p = self.doc.add_paragraph()
                        p.paragraph_format.left_indent = Inches(indent_level)
                        self._add_pill_badge(p, key.upper(), "17A2B8")
                    p.add_run("Removed options:")
                    for item in val['removed']:
                        p_sub = self.doc.add_paragraph()
                        p_sub.paragraph_format.left_indent = Inches(indent_level + 0.25)
                        self._add_pill_badge(p_sub, "REMOVED", "DC3545")
                        if isinstance(item, dict):
                             p_sub.add_run(self._format_schema_summary(item))
                        else:
                             p_sub.add_run(str(item))

            elif isinstance(val, dict) and 'old' in val and 'new' in val:
                p.add_run(f"Changed from '{val['old']}' to '{val['new']}'")
            elif isinstance(val, dict) and 'old_count' in val:
                p.add_run(f"Count changed from {val['old_count']} to {val['new_count']}")
            else:
                p.add_run(str(val))

    def _is_substantial_modification(self, item_name, changes, renamed_map, c_type='schemas'):
        """
        Determines if a modification is 'substantial' or just a rename-induced ref change.
        Returns True if substantial, False otherwise.
        """
        import copy
        filtered_changes = copy.deepcopy(changes)
        
        if c_type == 'schemas' and 'properties' in filtered_changes and 'modified' in filtered_changes['properties']:
            props_mod = filtered_changes['properties']['modified']
            props_to_remove = []
            
            for prop, p_diff in props_mod.items():
                # Check for $ref change
                ref_change = None
                if '$ref' in p_diff:
                    ref_change = p_diff['$ref']
                elif 'items' in p_diff and '$ref' in p_diff['items']: # Array of refs
                        ref_change = p_diff['items']['$ref']
                
                if ref_change:
                    old_ref = str(ref_change.get('old') or '')
                    new_ref = str(ref_change.get('new') or '')
                    old_simple = old_ref.split('/')[-1]
                    new_simple = new_ref.split('/')[-1]
                    
                    # Check if this is a known rename
                    if old_simple in renamed_map and renamed_map[old_simple] == new_simple:
                        props_to_remove.append(prop)
            
            for p in props_to_remove:
                del props_mod[p]
            
            if not props_mod:
                del filtered_changes['properties']['modified']
                if not filtered_changes['properties']:
                    del filtered_changes['properties']
        
        # Check if anything else remains (ignoring __rename_info__)
        keys = [k for k in filtered_changes.keys() if k != '__rename_info__']
        return len(keys) > 0

    def _add_components(self):
        self.doc.add_heading('Components', 1)
        
        # Order of presentation
        comp_types = ['schemas', 'parameters', 'responses', 'headers', 'securitySchemes', 'links', 'callbacks', 'examples']
        
        for c_type in comp_types:
            # Check if any changes exist for this type
            new_items = self.diff.new_components.get(c_type, [])
            rem_items = self.diff.removed_components.get(c_type, [])
            mod_items = self.diff.modified_components.get(c_type, {})
            
            # Renamed (Strict Partition)
            if self.diff.renamed_components.get(c_type):
                all_renames = self.diff.renamed_components[c_type]
                
                # Filter: Show here if NOT modified OR if modification is NOT substantial
                items_to_show = {}
                for old, new in all_renames.items():
                    if old not in mod_items:
                        items_to_show[old] = new
                    elif not self._is_substantial_modification(old, mod_items[old], all_renames, c_type):
                        items_to_show[old] = new
                
                if items_to_show:
                    self.doc.add_heading(f"Renamed {c_type}", 3)
                    for old_name in sorted(items_to_show.keys()):
                        new_name = items_to_show[old_name]
                        
                        # Badge in separate paragraph
                        p_badge = self.doc.add_paragraph()
                        p_badge.paragraph_format.left_indent = Inches(0.25)
                        p_badge.paragraph_format.space_after = Pt(0)
                        self._add_pill_badge(p_badge, "RENAMED", "17A2B8")
                        
                        # Name in Heading 4 (for Nav Pane) but styled as text
                        p = self.doc.add_heading('', level=4)
                        p.paragraph_format.left_indent = Inches(0.25)
                        p.paragraph_format.space_before = Pt(0)
                        
                        # Run 1: New Name (Bold)
                        r1 = p.add_run(new_name)
                        r1.font.color.rgb = RGBColor(0, 0, 0)
                        r1.font.bold = True
                        
                        # Run 2: Old Name (Normal)
                        r2 = p.add_run(f" (was {old_name})")
                        r2.font.color.rgb = RGBColor(0, 0, 0)
                        r2.font.bold = False
                    self.doc.add_paragraph().paragraph_format.space_after = Pt(12)

            # New
            if new_items:
                self.doc.add_heading(f"New {c_type}", 3)
                for item in sorted(new_items):
                    if c_type == 'schemas':
                        p_badge = self.doc.add_paragraph()
                        p_badge.paragraph_format.left_indent = Inches(0.25)
                        p_badge.paragraph_format.space_after = Pt(0)
                        self._add_pill_badge(p_badge, "NEW", "28A745")
                        
                        p = self.doc.add_heading(item, level=4)
                        p.paragraph_format.left_indent = Inches(0.25)
                        p.paragraph_format.space_before = Pt(0)
                        for run in p.runs:
                            run.font.color.rgb = RGBColor(0, 0, 0)
                            run.font.bold = False
                    else:
                        p = self.doc.add_paragraph()
                        p.paragraph_format.left_indent = Inches(0.25)
                        self._add_pill_badge(p, "NEW", "28A745")
                        p.add_run(f" {c_type[:-1].capitalize()}: {item}")
                self.doc.add_paragraph().paragraph_format.space_after = Pt(12)
                
            # Removed
            if rem_items:
                self.doc.add_heading(f"Removed {c_type}", 3)
                for item in sorted(rem_items):
                    if c_type == 'schemas':
                        p_badge = self.doc.add_paragraph()
                        p_badge.paragraph_format.left_indent = Inches(0.25)
                        p_badge.paragraph_format.space_after = Pt(0)
                        self._add_pill_badge(p_badge, "REMOVED", "DC3545")
                        
                        p = self.doc.add_heading(item, level=4)
                        p.paragraph_format.left_indent = Inches(0.25)
                        p.paragraph_format.space_before = Pt(0)
                        for run in p.runs:
                            run.font.color.rgb = RGBColor(0, 0, 0)
                            run.font.bold = False
                    else:
                        p = self.doc.add_paragraph()
                        p.paragraph_format.left_indent = Inches(0.25)
                        self._add_pill_badge(p, "REMOVED", "DC3545")
                        p.add_run(f" {c_type[:-1].capitalize()}: {item}")
                self.doc.add_paragraph().paragraph_format.space_after = Pt(12)
                
            # Modified (Strict Partition)
            if mod_items:
                # Filter: Show here ONLY if substantial modification
                filtered_mod_items = {}
                renamed_map = self.diff.renamed_components.get(c_type, {})
                
                for item_name, changes in mod_items.items():
                    # Use the helper to check substantiality
                    if self._is_substantial_modification(item_name, changes, renamed_map, c_type):
                        filtered_mod_items[item_name] = changes

                if filtered_mod_items:
                    self.doc.add_heading(f"Modified {c_type}", 3)
                    for item_name in sorted(filtered_mod_items.keys()):
                        changes = filtered_mod_items[item_name]
                        # Check if this item was also renamed
                        display_name = item_name
                        rename_note = ""
                        new_name = renamed_map.get(item_name)
                        if new_name:
                            display_name = new_name
                            rename_note = f" (was {item_name})"

                        # Use Pill Badge style instead of Heading 4 (but use Heading 4 for Nav Pane if Schema)
                        if c_type == 'schemas':
                            p_badge = self.doc.add_paragraph()
                            p_badge.paragraph_format.space_before = Pt(12)
                            p_badge.paragraph_format.left_indent = Inches(0.25)
                            p_badge.paragraph_format.space_after = Pt(0)
                            
                            if rename_note:
                                self._add_pill_badge(p_badge, "RENAMED & MODIFIED", "FFC107")
                            else:
                                self._add_pill_badge(p_badge, "MODIFIED", "FFC107")
                                
                            p = self.doc.add_heading('', level=4)
                            p.paragraph_format.left_indent = Inches(0.25)
                            p.paragraph_format.space_before = Pt(0)
                            
                            # Run 1: Display Name (Bold, Larger, Dark Blue)
                            r1 = p.add_run(display_name)
                            r1.font.color.rgb = RGBColor(0, 51, 102) # Dark Blue
                            r1.font.size = Pt(12) # Larger
                            r1.font.bold = True
                            
                            # Run 2: Rename Note (Normal)
                            if rename_note:
                                r2 = p.add_run(rename_note)
                                r2.font.color.rgb = RGBColor(0, 0, 0)
                                r2.font.bold = False

                        else:
                            p = self.doc.add_paragraph()
                            p.paragraph_format.space_before = Pt(12)
                            p.paragraph_format.left_indent = Inches(0.25)
                            
                            if rename_note:
                                self._add_pill_badge(p, "RENAMED & MODIFIED", "FFC107")
                            else:
                                self._add_pill_badge(p, "MODIFIED", "FFC107")
                            
                            p.add_run(display_name + rename_note) # No Bold
                        
                        # Special handling for Schemas properties table
                        if c_type == 'schemas' and 'properties' in changes and 'modified' in changes['properties']:
                            p_prop = self.doc.add_paragraph('Property Changes:')
                            p_prop.paragraph_format.left_indent = Inches(0.5) # Indent label
                            p_prop.paragraph_format.space_before = Pt(12) # Standard separation
                            p_prop.paragraph_format.space_after = Pt(4)
                            
                            # Calculate available width (Total 6.4 - 0.5 indent = 5.9)
                            total_avail = 6.4 - 0.5
                            col1 = 1.3
                            col2 = 0.7
                            col34 = (total_avail - col1 - col2) / 2
                            widths = [Inches(col1), Inches(col2), Inches(col34), Inches(col34)]
                            
                            table = self._create_table(4, widths)
                            tblPr = table._tblPr
                            tblInd = get_or_add_child(tblPr, 'w:tblInd', TBL_PR_ORDER)
                            tblInd.set(qn('w:w'), str(int(Inches(0.5).twips))) # Indent table
                            tblInd.set(qn('w:type'), 'dxa')
                            
                            self._style_header_row(table.rows[0], ['Property', 'Change', 'Old', 'New'])
                            
                            for prop, p_diff in changes['properties']['modified'].items():
                                row = table.add_row()
                                row.cells[0].text = prop
                                row.cells[1].text = 'Mod'
                                
                                cell_old = row.cells[2]
                                cell_new = row.cells[3]
                                p_old = cell_old.paragraphs[0]
                                p_new = cell_new.paragraphs[0]
                                
                                for k, v in p_diff.items():
                                    if k == 'description' and isinstance(v, dict) and 'old' in v:
                                        p_old.add_run("description: ")
                                        p_new.add_run("description: ")
                                        self._render_rich_diff(p_old, p_new, v['old'], v['new'])
                                        p_old.add_run("\n")
                                        p_new.add_run("\n")
                                    elif isinstance(v, dict) and 'old' in v:
                                        p_old.add_run(f"{k}: {v['old']}\n")
                                        p_new.add_run(f"{k}: {v['new']}\n")
                                    else:
                                        p_old.add_run(f"{k}: (complex)\n")
                                        p_new.add_run(f"{k}: (complex)\n")
                                
                                for cell in row.cells:
                                    self._style_body_cell(cell)
                            self.doc.add_paragraph().paragraph_format.space_after = Pt(12)
                    
                        # Track if we printed anything for this item
                        content_printed = False

                        if 'properties' in changes:
                            if 'new' in changes['properties'] and changes['properties']['new']:
                                p_new = self.doc.add_paragraph("New Properties:")
                                p_new.paragraph_format.left_indent = Inches(0.5) # Indented under schema
                                p_new.paragraph_format.space_before = Pt(12) # Standard separation
                                p_new.paragraph_format.space_after = Pt(4)
                                
                                for prop in changes['properties']['new']:
                                    p = self.doc.add_paragraph()
                                    p.paragraph_format.left_indent = Inches(0.75) # Further indented
                                    p.paragraph_format.space_after = Pt(2)
                                    self._add_pill_badge(p, "NEW PROP", "28A745")
                                    p.add_run(prop)
                                content_printed = True

                            if 'removed' in changes['properties'] and changes['properties']['removed']:
                                p_rem = self.doc.add_paragraph("Removed Properties:")
                                p_rem.paragraph_format.left_indent = Inches(0.5) # Indented under schema
                                p_rem.paragraph_format.space_before = Pt(12) # Standard separation
                                p_rem.paragraph_format.space_after = Pt(4)
                                
                                for prop in changes['properties']['removed']:
                                    p = self.doc.add_paragraph()
                                    p.paragraph_format.left_indent = Inches(0.75) # Further indented
                                    p.paragraph_format.space_after = Pt(2)
                                    self._add_pill_badge(p, "REMOVED PROP", "DC3545")
                                    p.add_run(prop)
                                content_printed = True
                            
                            if 'modified' in changes['properties']:
                                content_printed = True # Table already added above

                        # 5. Attribute Changes Table
                        # Collect simple attribute changes (not properties, ref, or combinators)
                        attr_changes = {}
                        ignored_keys = ['properties', '$ref', 'oneOf', 'allOf', 'anyOf', '__rename_info__']
                        
                        for key, val in changes.items():
                            if key in ignored_keys: continue
                            if isinstance(val, dict) and 'old' in val and 'new' in val:
                                attr_changes[key] = val
                        
                        if attr_changes:
                            content_printed = True
                            p_attr = self.doc.add_paragraph('Attribute Changes:')
                            p_attr.paragraph_format.left_indent = Inches(0.5)
                            p_attr.paragraph_format.space_before = Pt(12)
                            p_attr.paragraph_format.space_after = Pt(4)
                            
                            # Calculate available width (7.0 - 0.5 indent = 6.5)
                            widths = [Inches(1.5), Inches(0.8), Inches(2.1), Inches(2.1)]
                            table = self._create_table(4, widths)
                            tblPr = table._tblPr
                            tblInd = get_or_add_child(tblPr, 'w:tblInd', TBL_PR_ORDER)
                            tblInd.set(qn('w:w'), str(int(Inches(0.5).twips)))
                            tblInd.set(qn('w:type'), 'dxa')
                            
                            self._style_header_row(table.rows[0], ['Attribute', 'Change', 'Old Value', 'New Value'])
                            
                            for attr, val in attr_changes.items():
                                row = table.add_row()
                                row.cells[0].text = attr
                                row.cells[1].text = 'Mod'
                                
                                if attr == 'description':
                                    self._render_rich_diff(row.cells[2].paragraphs[0], row.cells[3].paragraphs[0], val['old'], val['new'])
                                else:
                                    old_val = val['old']
                                    new_val = val['new']
                                    if isinstance(old_val, (dict, list)):
                                        import json
                                        old_val = json.dumps(old_val, indent=2)
                                    if isinstance(new_val, (dict, list)):
                                        import json
                                        new_val = json.dumps(new_val, indent=2)
                                    
                                    row.cells[2].text = str(old_val)
                                    row.cells[3].text = str(new_val)
                                
                                for cell in row.cells:
                                    self._style_body_cell(cell)
                            
                            self.doc.add_paragraph().paragraph_format.space_after = Pt(6)

                        for key, val in changes.items():
                            if key in ['properties', '__rename_info__']: continue
                            # Skip if already handled in Attribute Changes table
                            if key in attr_changes: continue
                            
                            content_printed = True
                            p = self.doc.add_paragraph()
                            p.paragraph_format.left_indent = Inches(0.5)
                            
                            if key in ['oneOf', 'allOf', 'anyOf']:
                                # Keyword Style (Bold Monospace, Dark Blue)
                                run = p.add_run(key)
                                run.font.name = 'Consolas'
                                run.font.bold = True
                                run.font.color.rgb = RGBColor(0, 51, 102) # Dark Blue
                                p.add_run(" ") # Spacer
                            else:
                                self._add_pill_badge(p, key.upper(), "17A2B8")
                            if isinstance(val, dict) and ('added' in val or 'removed' in val):
                                added_items = val.get('added', [])
                                removed_items = val.get('removed', [])
                                
                                # Check for rename pairs in combinators
                                renamed_pairs = []
                                indices_added = []
                                indices_removed = []
                                
                                if key in ['oneOf', 'allOf', 'anyOf'] and c_type == 'schemas':
                                    renamed_map = self.diff.renamed_components.get('schemas', {})
                                    for i_rem, rem_item in enumerate(removed_items):
                                        rem_ref = None
                                        if isinstance(rem_item, str): rem_ref = rem_item
                                        elif isinstance(rem_item, dict) and '$ref' in rem_item: rem_ref = rem_item['$ref']
                                        
                                        if rem_ref:
                                            rem_simple = rem_ref.split('/')[-1]
                                            if rem_simple in renamed_map:
                                                new_simple = renamed_map[rem_simple]
                                                for i_add, add_item in enumerate(added_items):
                                                    if i_add in indices_added: continue
                                                    add_ref = None
                                                    if isinstance(add_item, str): add_ref = add_item
                                                    elif isinstance(add_item, dict) and '$ref' in add_item: add_ref = add_item['$ref']
                                                    
                                                    if add_ref and add_ref.split('/')[-1] == new_simple:
                                                        renamed_pairs.append((rem_ref, add_ref))
                                                        indices_removed.append(i_rem)
                                                        indices_added.append(i_add)
                                                        break
                                
                                # Display Renamed Pairs
                                if renamed_pairs:
                                    p.add_run("Ref changes:")
                                    for old_ref, new_ref in renamed_pairs:
                                        p_sub = self.doc.add_paragraph()
                                        p_sub.paragraph_format.left_indent = Inches(0.75)
                                        self._add_pill_badge(p_sub, "REF RENAMED", "FFC107")
                                        p_sub.add_run(f"'{old_ref}' \u2192 '{new_ref}'")

                                # Display remaining Added
                                remaining_added = [item for i, item in enumerate(added_items) if i not in indices_added]
                                if remaining_added:
                                    if renamed_pairs: 
                                        p_lbl = self.doc.add_paragraph()
                                        p_lbl.paragraph_format.left_indent = Inches(0.5)
                                        self._add_pill_badge(p_lbl, key.upper(), "17A2B8")
                                    p.add_run("Added options:")
                                    for item in remaining_added:
                                        p_sub = self.doc.add_paragraph()
                                        p_sub.paragraph_format.left_indent = Inches(0.75)
                                        self._add_pill_badge(p_sub, "ADDED", "28A745")
                                        if c_type == 'schemas':
                                            p_sub.add_run(self._format_schema_summary(item))
                                        else:
                                            p_sub.add_run(str(item))

                                # Display remaining Removed
                                remaining_removed = [item for i, item in enumerate(removed_items) if i not in indices_removed]
                                if remaining_removed:
                                    if remaining_added or renamed_pairs:
                                        p = self.doc.add_paragraph()
                                        p.paragraph_format.left_indent = Inches(0.5)
                                        self._add_pill_badge(p, key.upper(), "17A2B8")
                                    p.add_run("Removed options:")
                                    for item in remaining_removed:
                                        p_sub = self.doc.add_paragraph()
                                        p_sub.paragraph_format.left_indent = Inches(0.75)
                                        self._add_pill_badge(p_sub, "REMOVED", "DC3545")
                                        if c_type == 'schemas':
                                            p_sub.add_run(self._format_schema_summary(item))
                                        else:
                                            p_sub.add_run(str(item))

                            elif isinstance(val, dict) and 'old' in val and 'new' in val:
                                 p.add_run(f" - {key} changed:")
                                 # Use table for generic value change
                                 widths = [Inches(3.0), Inches(3.0)]
                                 table = self._create_table(2, widths)
                                 self._style_header_row(table.rows[0], ['Old Value', 'New Value'])
                                 row = table.add_row()
                                 
                                 # Format values
                                 old_val = val['old']
                                 new_val = val['new']
                                 if isinstance(old_val, (dict, list)):
                                     import json
                                     old_val = json.dumps(old_val, indent=2)
                                 if isinstance(new_val, (dict, list)):
                                     import json
                                     new_val = json.dumps(new_val, indent=2)
                                     
                                 row.cells[0].text = str(old_val)
                                 row.cells[1].text = str(new_val)
                                 for cell in row.cells:
                                     self._style_body_cell(cell)
                                 self.doc.add_paragraph().paragraph_format.space_after = Pt(6)
                            else:
                                 # Fallback
                                 p.add_run(str(val))
                        
                        if not content_printed:
                            p = self.doc.add_paragraph()
                            p.paragraph_format.left_indent = Inches(0.5)
                            p.add_run("Metadata or internal structure modified.")
                            p.font.italic = True
                            # Debug info
                            p.add_run(f" (Keys: {', '.join(changes.keys())})")

                        # --- AFFECTED ENDPOINTS SECTION ---
                        # Look up impacts using the NEW name (as that's what's in spec2)
                        # We place this LAST as requested.
                        if c_type == 'schemas':
                            impact_name = item_name
                            renamed_map = self.diff.renamed_components.get('schemas', {})
                            if item_name in renamed_map:
                                impact_name = renamed_map[item_name]
                                
                            impacts = self.tracer.get_impacted_endpoints(impact_name)
                            
                            if impacts:
                                p_impact = self.doc.add_paragraph('Affected Endpoints:')
                                p_impact.paragraph_format.left_indent = Inches(0.5)
                                p_impact.paragraph_format.space_before = Pt(12) # Standard separation
                                p_impact.paragraph_format.space_after = Pt(4)
                                
                                # Table for impacts
                                # Widths: Method(0.8), Path(3.0), Context(2.7) -> Total 6.5
                                widths = [Inches(0.8), Inches(3.0), Inches(2.7)]
                                table = self._create_table(3, widths)
                                tblPr = table._tblPr
                                tblInd = get_or_add_child(tblPr, 'w:tblInd', TBL_PR_ORDER)
                                tblInd.set(qn('w:w'), str(int(Inches(0.5).twips)))
                                tblInd.set(qn('w:type'), 'dxa')
                                
                                self._style_header_row(table.rows[0], ['Method', 'Endpoint', 'Context'])
                                
                                sorted_impacts = sorted(impacts, key=lambda x: (x['path'], x['method']))
                                for impact in sorted_impacts:
                                    row = table.add_row()
                                    row.cells[0].text = impact['method']
                                    row.cells[1].text = impact['path']
                                    row.cells[2].text = impact['context']
                                    
                                    # Style Method (Bold)
                                    row.cells[0].paragraphs[0].runs[0].font.bold = True
                                    
                                    for cell in row.cells:
                                        self._style_body_cell(cell)
                                
                                self.doc.add_paragraph().paragraph_format.space_after = Pt(8)
                        # ----------------------------------

    def _render_rich_diff(self, p_old, p_new, text_old, text_new):
        """Renders character-level diff with shading (50% lighter colors) and robust opcode merging."""
        if not isinstance(text_old, str): text_old = str(text_old or "")
        if not isinstance(text_new, str): text_new = str(text_new or "")
        
        s = difflib.SequenceMatcher(None, text_old, text_new)
        opcodes = s.get_opcodes()
        
        # Robust Merging: combine consecutive (delete|replace|insert) into one big replace
        # unless they are separated by 'equal'.
        merged_opcodes = []
        i = 0
        while i < len(opcodes):
            tag, i1, i2, j1, j2 = opcodes[i]
            if tag == 'equal':
                merged_opcodes.append((tag, i1, i2, j1, j2))
                i += 1
            else:
                # Start merging a block of changes
                curr_tag = 'replace' # default if multiple mixed
                start_i1, last_i2 = i1, i2
                start_j1, last_j2 = j1, j2
                
                # If it's pure delete or pure insert first, we might keep it, 
                # but user wants 'yellow' for most context changes.
                
                k = i + 1
                while k < len(opcodes) and opcodes[k][0] != 'equal':
                    next_tag, ni1, ni2, nj1, nj2 = opcodes[k]
                    last_i2 = ni2
                    last_j2 = nj2
                    k += 1
                
                # If we only have ONE op and it was delete or insert, we can keep the tag.
                # But if we combined multiple ops, it's definitely a 'replace'.
                if k == i + 1:
                    merged_opcodes.append((tag, i1, i2, j1, j2))
                else:
                    merged_opcodes.append(('replace', start_i1, last_i2, start_j1, last_j2))
                i = k

    def _render_rich_diff(self, p_old, p_new, text_old, text_new):
        """Renders description diff with paragraph-level tracking and splitting lopsided replacements."""
        if not isinstance(text_old, str): text_old = str(text_old or "")
        if not isinstance(text_new, str): text_new = str(text_new or "")
        
        # 1. Split into paragraphs
        lines_old = text_old.splitlines(keepends=True)
        lines_new = text_new.splitlines(keepends=True)
        
        s = difflib.SequenceMatcher(None, lines_old, lines_new, autojunk=False)
        opcodes = s.get_opcodes()
        
        # Split "lopsided" replace opcodes: e.g. replace 1..4 with 1 -> replace 1 with 1 + delete 2..4
        refined_opcodes = []
        for tag, i1, i2, j1, j2 in opcodes:
            if tag == 'replace':
                old_count = i2 - i1
                new_count = j2 - j1
                if old_count > new_count:
                    refined_opcodes.append(('replace', i1, i1 + new_count, j1, j2))
                    refined_opcodes.append(('delete', i1 + new_count, i2, j2, j2))
                elif new_count > old_count:
                    refined_opcodes.append(('replace', i1, i2, j1, j1 + old_count))
                    refined_opcodes.append(('insert', i2, i2, j1 + old_count, j2))
                else:
                    refined_opcodes.append((tag, i1, i2, j1, j2))
            else:
                refined_opcodes.append((tag, i1, i2, j1, j2))

        def apply_shading(run, color_hex):
            rPr = run._r.get_or_add_rPr()
            shd = get_or_add_child(rPr, 'w:shd', R_PR_ORDER)
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:fill'), color_hex)

        def render_word_diff(para_old, para_new, t_old, t_new):
            import re
            def split_words(text):
                return re.findall(r'\w+|[^\w\s]|\s+', text)
            
            w_old = split_words(t_old)
            w_new = split_words(t_new)
            ws = difflib.SequenceMatcher(None, w_old, w_new, autojunk=False)
            w_ops = ws.get_opcodes()
            
            # Refined merging: bridge only very small gaps (e.g. whitespace)
            # to keep keywords distinct if they are pure removals.
            merged_w_ops = []
            wi = 0
            while wi < len(w_ops):
                tag, i1, i2, j1, j2 = w_ops[wi]
                if tag == 'equal':
                    merged_w_ops.append((tag, i1, i2, j1, j2))
                    wi += 1
                else:
                    sw_i1, lw_i2 = i1, i2
                    sw_j1, lw_j2 = j1, j2
                    actual_o = tag in ('delete', 'replace')
                    actual_n = tag in ('insert', 'replace')
                    
                    wk = wi + 1
                    while wk < len(w_ops):
                        nt, ni1, ni2, nj1, nj2 = w_ops[wk]
                        is_bridge = False
                        if nt == 'equal':
                            eq_text = "".join(w_old[ni1:ni2])
                            if len(eq_text) <= 3 and '\n' not in eq_text: is_bridge = True
                        
                        if nt != 'equal' or is_bridge:
                            if nt in ('delete', 'replace'): actual_o = True
                            if nt in ('insert', 'replace'): actual_n = True
                            lw_i2, lw_j2 = ni2, nj2
                            wk += 1
                        else: break
                    
                    wt = 'replace' if (actual_o and actual_n) else ('delete' if actual_o else 'insert')
                    merged_w_ops.append((wt, sw_i1, lw_i2, sw_j1, lw_j2))
                    wi = wk

            for wt, wi1, wi2, wj1, wj2 in merged_w_ops:
                txt_o = "".join(w_old[wi1:wi2])
                txt_n = "".join(w_new[wj1:wj2])
                if wt == 'equal':
                    para_old.add_run(txt_o)
                    para_new.add_run(txt_n)
                elif wt == 'replace':
                    # MODIFICA: Giallo in entrambi
                    apply_shading(para_old.add_run(txt_o), "FFF3CD")
                    apply_shading(para_new.add_run(txt_n), "FFF3CD")
                elif wt == 'delete':
                    # RIMOZIONE: Rosso in Old
                    apply_shading(para_old.add_run(txt_o), "F8D7DA")
                elif wt == 'insert':
                    # AGGIUNTA: Verde in New
                    apply_shading(para_new.add_run(txt_n), "D4EDDA")

        for tag, i1, i2, j1, j2 in refined_opcodes:
            if tag == 'equal':
                p_old.add_run("".join(lines_old[i1:i2]))
                p_new.add_run("".join(lines_new[j1:j2]))
            elif tag == 'delete':
                apply_shading(p_old.add_run("".join(lines_old[i1:i2])), "F8D7DA")
            elif tag == 'insert':
                apply_shading(p_new.add_run("".join(lines_new[j1:j2])), "D4EDDA")
            elif tag == 'replace':
                render_word_diff(p_old, p_new, "".join(lines_old[i1:i2]), "".join(lines_new[j1:j2]))

    def _render_rich_diff_inline(self, p, text_old, text_new):
        """Renders description diff inline with line splitting and color accuracy."""
        if not isinstance(text_old, str): text_old = str(text_old or "")
        if not isinstance(text_new, str): text_new = str(text_new or "")
        
        lines_old = text_old.splitlines(keepends=True)
        lines_new = text_new.splitlines(keepends=True)
        
        s = difflib.SequenceMatcher(None, lines_old, lines_new, autojunk=False)
        opcodes = s.get_opcodes()
        
        refined_opcodes = []
        for tag, i1, i2, j1, j2 in opcodes:
            if tag == 'replace':
                old_count = i2 - i1
                new_count = j2 - j1
                if old_count > new_count:
                    refined_opcodes.append(('replace', i1, i1 + new_count, j1, j2))
                    refined_opcodes.append(('delete', i1 + new_count, i2, j2, j2))
                elif new_count > old_count:
                    refined_opcodes.append(('replace', i1, i2, j1, j1 + old_count))
                    refined_opcodes.append(('insert', i2, i2, j1 + old_count, j2))
                else: refined_opcodes.append((tag, i1, i2, j1, j2))
            else: refined_opcodes.append((tag, i1, i2, j1, j2))

        def apply_shading(run, color_hex):
            rPr = run._r.get_or_add_rPr()
            shd = get_or_add_child(rPr, 'w:shd', R_PR_ORDER)
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:fill'), color_hex)

        def render_word_diff_inline(para, t_old, t_new):
            import re
            def split_words(text):
                return re.findall(r'\w+|[^\w\s]|\s+', text)
            
            w_old = split_words(t_old)
            w_new = split_words(t_new)
            ws = difflib.SequenceMatcher(None, w_old, w_new, autojunk=False)
            w_ops = ws.get_opcodes()
            
            merged_w_ops = []
            wi = 0
            while wi < len(w_ops):
                tag, i1, i2, j1, j2 = w_ops[wi]
                if tag == 'equal':
                    merged_w_ops.append((tag, i1, i2, j1, j2))
                    wi += 1
                else:
                    sw_i1, lw_i2 = i1, i2
                    sw_j1, lw_j2 = j1, j2
                    actual_o = tag in ('delete', 'replace')
                    actual_n = tag in ('insert', 'replace')
                    wk = wi + 1
                    while wk < len(w_ops):
                        nt, ni1, ni2, nj1, nj2 = w_ops[wk]
                        is_bridge = False
                        if nt == 'equal':
                            eq_text = "".join(w_old[ni1:ni2])
                            if len(eq_text) <= 3 and '\n' not in eq_text: is_bridge = True
                        if nt != 'equal' or is_bridge:
                            if nt in ('delete', 'replace'): actual_o = True
                            if nt in ('insert', 'replace'): actual_n = True
                            lw_i2, lw_j2 = ni2, nj2
                            wk += 1
                        else: break
                    wt = 'replace' if (actual_o and actual_n) else ('delete' if actual_o else 'insert')
                    merged_w_ops.append((wt, sw_i1, lw_i2, sw_j1, lw_j2))
                    wi = wk

            for wt, wi1, wi2, wj1, wj2 in merged_w_ops:
                txt = "".join(w_old[wi1:wi2])
                if wt == 'equal': para.add_run(txt)
                elif wt == 'delete': apply_shading(para.add_run(txt), "F8D7DA")
                elif wt == 'replace': apply_shading(para.add_run(txt), "FFF3CD")
            para.add_run(" \u2192 ")
            for wt, wi1, wi2, wj1, wj2 in merged_w_ops:
                txt = "".join(w_new[wj1:wj2])
                if wt == 'equal': para.add_run(txt)
                elif wt == 'insert': apply_shading(para.add_run(txt), "D4EDDA")
                elif wt == 'replace': apply_shading(para.add_run(txt), "FFF3CD")

        for tag, i1, i2, j1, j2 in refined_opcodes:
            txt_o = "".join(lines_old[i1:i2])
            txt_n = "".join(lines_new[j1:j2])
            if tag == 'equal': p.add_run(txt_o)
            elif tag == 'delete': apply_shading(p.add_run(txt_o), "F8D7DA")
            elif tag == 'insert': 
                p.add_run(" [+] ")
                apply_shading(p.add_run(txt_n), "D4EDDA")
            elif tag == 'replace':
                render_word_diff_inline(p, txt_o, txt_n)

    def _format_schema_summary(self, schema):
        if '$ref' in schema:
            return schema['$ref']
        elif 'type' in schema:
            return f"Type: {schema['type']}"
        return "Complex Schema Object"
