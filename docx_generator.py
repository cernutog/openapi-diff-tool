from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

# --- OXML Helpers (Safe Insertion) ---
def get_or_add_child(parent, tag_name, order_list):
    child = parent.find(qn(tag_name))
    if child is not None:
        return child
    child = OxmlElement(tag_name)
    try:
        my_idx = order_list.index(tag_name)
    except ValueError:
        parent.append(child)
        return child
    for i in range(my_idx + 1, len(order_list)):
        next_tag = order_list[i]
        next_element = parent.find(qn(next_tag))
        if next_element is not None:
            parent.insert(parent.index(next_element), child)
            return child
    parent.append(child)
    return child

TBL_PR_ORDER = ['w:tblStyle', 'w:tblpPr', 'w:tblOverlap', 'w:bidiVisual', 'w:tblStyleRowBandSize', 'w:tblStyleColBandSize', 'w:tblW', 'w:jc', 'w:tblCellSpacing', 'w:tblInd', 'w:tblBorders', 'w:shd', 'w:tblLayout', 'w:tblCellMar', 'w:tblLook']
TC_PR_ORDER = ['w:cnfStyle', 'w:tcW', 'w:gridSpan', 'w:hMerge', 'w:vMerge', 'w:tcBorders', 'w:shd', 'w:noWrap', 'w:tcMar', 'w:textDirection', 'w:tcFitText', 'w:vAlign', 'w:hideMark', 'w:headers', 'w:cellIns', 'w:cellDel', 'w:cellMerge', 'w:tcPrChange']
P_PR_ORDER = ['w:pStyle', 'w:keepNext', 'w:keepLines', 'w:pageBreakBefore', 'w:framePr', 'w:widowControl', 'w:numPr', 'w:suppressLineNumbers', 'w:pBdr', 'w:shd', 'w:tabs', 'w:suppressAutoHyphens', 'w:kinsoku', 'w:wordWrap', 'w:overflowPunct', 'w:topLinePunct', 'w:autoSpaceDE', 'w:autoSpaceDN', 'w:bidi', 'w:adjustRightInd', 'w:snapToGrid', 'w:spacing', 'w:ind', 'w:contextualSpacing', 'w:mirrorIndents', 'w:suppressOverlap', 'w:jc', 'w:textDirection', 'w:textAlignment', 'w:textboxTightWrap', 'w:outlineLvl', 'w:divId', 'w:cnfStyle', 'w:rPr', 'w:sectPr', 'w:pPrChange']
R_PR_ORDER = ['w:rStyle', 'w:rFonts', 'w:b', 'w:bCs', 'w:i', 'w:iCs', 'w:caps', 'w:smallCaps', 'w:strike', 'w:dstrike', 'w:outline', 'w:shadow', 'w:emboss', 'w:imprint', 'w:noProof', 'w:snapToGrid', 'w:vanish', 'w:webHidden', 'w:color', 'w:spacing', 'w:w', 'w:kern', 'w:position', 'w:sz', 'w:szCs', 'w:highlight', 'w:u', 'w:effect', 'w:bdr', 'w:shd', 'w:fitText', 'w:vertAlign', 'w:rtl', 'w:cs', 'w:em', 'w:lang', 'w:eastAsianLayout', 'w:specVanish', 'w:oMath']

class DocxReportGenerator:
    def __init__(self, diff):
        self.diff = diff
        self.doc = Document()
        self._setup_styles()
        self._setup_page_layout()
        self._add_header_footer()

    def _setup_styles(self):
        # Normal
        style = self.doc.styles['Normal']
        style.font.name = 'Segoe UI'
        style.font.size = Pt(10)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.15

        # Table Text
        if 'Table Text' not in self.doc.styles:
            s = self.doc.styles.add_style('Table Text', 1)
            s.font.name = 'Segoe UI'
            s.font.size = Pt(9)
            s.paragraph_format.space_after = Pt(2)
            s.paragraph_format.line_spacing = 1.0

        # Headings
        h1 = self.doc.styles['Heading 1']
        h1.font.name = 'Segoe UI'
        h1.font.size = Pt(16)
        h1.font.bold = True
        h1.font.color.rgb = RGBColor(0, 51, 102) # Dark Navy
        h1.paragraph_format.space_before = Pt(24)
        h1.paragraph_format.space_after = Pt(6)
        # Border for H1
        pPr = h1._element.get_or_add_pPr()
        pbdr = get_or_add_child(pPr, 'w:pBdr', P_PR_ORDER)
        pbdr.clear()
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6') # Thicker
        bottom.set(qn('w:space'), '4')
        bottom.set(qn('w:color'), '003366')
        pbdr.append(bottom)

        h2 = self.doc.styles['Heading 2']
        h2.font.name = 'Segoe UI'
        h2.font.size = Pt(13)
        h2.font.bold = True
        h2.font.color.rgb = RGBColor(0, 51, 102)
        h2.paragraph_format.space_before = Pt(18)
        h2.paragraph_format.space_after = Pt(6)
        # Border for H2 (lighter)
        pPr2 = h2._element.get_or_add_pPr()
        pbdr2 = get_or_add_child(pPr2, 'w:pBdr', P_PR_ORDER)
        pbdr2.clear()
        bottom2 = OxmlElement('w:bottom')
        bottom2.set(qn('w:val'), 'single')
        bottom2.set(qn('w:sz'), '4')
        bottom2.set(qn('w:space'), '1')
        bottom2.set(qn('w:color'), 'CCCCCC') # Light Grey
        pbdr2.append(bottom2)

        h3 = self.doc.styles['Heading 3']
        h3.font.name = 'Segoe UI'
        h3.font.size = Pt(11)
        h3.font.bold = True
        h3.font.color.rgb = RGBColor(51, 51, 51)

        # List Styles (Ensure they exist)
        for i in range(2, 6):
            style_name = f'List Bullet {i}'
            if style_name not in self.doc.styles:
                s = self.doc.styles.add_style(style_name, 1) # 1 = Paragraph Style
                s.base_style = self.doc.styles['List Bullet']
                s.font.name = 'Segoe UI'
                # Indentation (approximate)
                s.paragraph_format.left_indent = Inches(0.25 * i)
                s.paragraph_format.first_line_indent = Inches(-0.25)

    def _setup_page_layout(self):
        section = self.doc.sections[0]
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)

    def _add_header_footer(self):
        section = self.doc.sections[0]
        # Header
        h_para = section.header.paragraphs[0]
        h_para.text = "OpenAPI Comparison Report - Confidential"
        h_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        h_para.runs[0].font.size = Pt(8)
        h_para.runs[0].font.color.rgb = RGBColor(128, 128, 128)
        # Footer
        f_para = section.footer.paragraphs[0]
        f_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        f_para.text = f"Generated on {datetime.datetime.now().strftime('%Y-%m-%d')} | Page "
        f_para.runs[0].font.size = Pt(8)
        f_para.runs[0].font.color.rgb = RGBColor(128, 128, 128)
        # Page number
        self._add_page_number(f_para)

    def _add_page_number(self, paragraph):
        run = paragraph.add_run()
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        run._r.append(fldChar1)
        instrText = OxmlElement('w:instrText')
        instrText.text = "PAGE"
        run_instr = paragraph.add_run()
        run_instr._r.append(instrText)
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        run_end = paragraph.add_run()
        run_end._r.append(fldChar2)

    def _create_table(self, cols, widths):
        table = self.doc.add_table(rows=1, cols=cols)
        table.style = 'Table Grid'
        table.autofit = False
        
        # Set fixed layout and width
        tblPr = table._tblPr
        tblW = get_or_add_child(tblPr, 'w:tblW', TBL_PR_ORDER)
        tblW.set(qn('w:w'), str(int(Inches(7.0).twips)))
        tblW.set(qn('w:type'), 'dxa')
        
        tblLayout = get_or_add_child(tblPr, 'w:tblLayout', TBL_PR_ORDER)
        tblLayout.set(qn('w:type'), 'fixed')
        
        # Set Grid
        tblGrid = table._element.find(qn('w:tblGrid'))
        if tblGrid is None:
            tblGrid = OxmlElement('w:tblGrid')
            table._element.insert(table._element.index(tblPr) + 1, tblGrid)
        else:
            tblGrid.clear()
            
        for w in widths:
            col = OxmlElement('w:gridCol')
            col.set(qn('w:w'), str(int(w.twips)))
            tblGrid.append(col)
            
        # Set initial cell widths
        for i, cell in enumerate(table.rows[0].cells):
            cell.width = widths[i]
            
        return table

    def _style_header_row(self, row, headers):
        for i, text in enumerate(headers):
            cell = row.cells[i]
            cell.text = text
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.runs[0]
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            
            tcPr = cell._element.get_or_add_tcPr()
            shd = get_or_add_child(tcPr, 'w:shd', TC_PR_ORDER)
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), '003366')
            
            # Margins
            tcMar = get_or_add_child(tcPr, 'w:tcMar', TC_PR_ORDER)
            tcMar.clear()
            for tag, val in [('w:top', 100), ('w:start', 100), ('w:bottom', 100), ('w:end', 100)]:
                node = OxmlElement(tag)
                node.set(qn('w:w'), str(val))
                node.set(qn('w:type'), 'dxa')
                tcMar.append(node)

    def _style_body_cell(self, cell):
        for p in cell.paragraphs:
            p.style = 'Table Text'
        tcPr = cell._element.get_or_add_tcPr()
        tcMar = get_or_add_child(tcPr, 'w:tcMar', TC_PR_ORDER)
        tcMar.clear()
        for tag, val in [('w:top', 50), ('w:start', 100), ('w:bottom', 50), ('w:end', 100)]:
            node = OxmlElement(tag)
            node.set(qn('w:w'), str(val))
            node.set(qn('w:type'), 'dxa')
            tcMar.append(node)

    def _add_badge(self, paragraph, text, bg_color='EEEEEE', text_color='000000'):
        """Adds a text run with a background color (shading) to simulate a badge."""
        run = paragraph.add_run(f" {text} ")
        run.font.size = Pt(8)
        run.font.bold = True
        run.font.color.rgb = RGBColor(int(text_color[:2], 16), int(text_color[2:4], 16), int(text_color[4:], 16))
        
        rPr = run._element.get_or_add_rPr()
        shd = get_or_add_child(rPr, 'w:shd', R_PR_ORDER)
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), bg_color)
        
        # Add a small space after
        paragraph.add_run(" ")

    def generate(self, output_path):
        # Title
        t = self.doc.add_heading('OpenAPI Changes Report', 0)
        t.alignment = WD_ALIGN_PARAGRAPH.CENTER
        s = self.doc.add_paragraph('Executive Summary', style='Subtitle')
        s.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        self._add_dashboard()
        self._add_general_info()
        self._add_endpoints()
        self._add_components()
        
        self.doc.save(output_path)

    def _add_dashboard(self):
        self.doc.add_heading('Dashboard', 1)
        widths = [Inches(2.3), Inches(2.3), Inches(2.4)]
        table = self._create_table(3, widths)
        self._style_header_row(table.rows[0], ['New Endpoints', 'Removed Endpoints', 'Modified Schemas'])
        
        row = table.add_row()
        row.cells[0].text = str(len(self.diff.new_paths))
        row.cells[1].text = str(len(self.diff.removed_paths))
        row.cells[2].text = str(len(self.diff.modified_components.get('schemas', {})))
        
        for cell in row.cells:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.paragraphs[0].runs[0].font.size = Pt(14)
            cell.paragraphs[0].runs[0].font.bold = True
            self._style_body_cell(cell)

    def _add_general_info(self):
        self.doc.add_heading('General Info', 1)
        if not self.diff.info_changes:
            self.doc.add_paragraph('No changes in General Info.')
            return

        widths = [Inches(1.5), Inches(2.75), Inches(2.75)]
        table = self._create_table(3, widths)
        self._style_header_row(table.rows[0], ['Field', 'Old Value', 'New Value'])
        
        for key, change in self.diff.info_changes.items():
            row = table.add_row()
            row.cells[0].text = str(key)
            row.cells[1].text = str(change['old'])
            row.cells[2].text = str(change['new'])
            for cell in row.cells:
                self._style_body_cell(cell)

    def _add_endpoints(self):
        self.doc.add_heading('Endpoints Summary', 1)
        
        # New
        self.doc.add_heading('New Endpoints', 2)
        if self.diff.new_paths:
            for path in self.diff.new_paths:
                p = self.doc.add_paragraph(style='List Bullet')
                self._add_badge(p, "NEW", "28A745", "FFFFFF") # Green
                p.add_run(path)
        else:
            self.doc.add_paragraph('No new endpoints.')

        # Modified
        self.doc.add_heading('Modified Endpoints', 2)
        if self.diff.modified_paths:
            for path, changes in self.diff.modified_paths.items():
                self.doc.add_heading(path, 3)
                
                if 'new_ops' in changes:
                    p = self.doc.add_paragraph(style='List Bullet')
                    self._add_badge(p, "NEW METHOD", "17A2B8", "FFFFFF") # Cyan
                    p.add_run(f"{', '.join(changes['new_ops']).upper()}")
                
                if 'removed_ops' in changes:
                    p = self.doc.add_paragraph(style='List Bullet')
                    self._add_badge(p, "REMOVED METHOD", "DC3545", "FFFFFF") # Red
                    p.add_run(f"{', '.join(changes['removed_ops']).upper()}")
                
                if 'modified_ops' in changes:
                    p = self.doc.add_paragraph(style='List Bullet')
                    self._add_badge(p, "MODIFIED METHOD", "FFC107", "000000") # Yellow
                    p.add_run("The following methods have changed:")
                    
                    for op, op_changes in changes['modified_ops'].items():
                        p_op = self.doc.add_paragraph(style='List Bullet 2')
                        self._add_badge(p_op, op.upper(), "6C757D", "FFFFFF") # Grey
                        
                        # Parameters
                        if 'parameters' in op_changes:
                            self._add_parameter_changes(op_changes['parameters'])
                            
                        # Request Body
                        if 'requestBody' in op_changes:
                            self._add_request_body_changes(op_changes['requestBody'])
                            
                        # Responses
                        if 'responses' in op_changes:
                            self._add_response_changes(op_changes['responses'])

        else:
            self.doc.add_paragraph('No modified endpoints.')

    def _add_parameter_changes(self, params_diff):
        if 'new' in params_diff:
            for param in params_diff['new']:
                p = self.doc.add_paragraph(style='List Bullet 3')
                self._add_badge(p, "NEW PARAM", "28A745", "FFFFFF")
                p.add_run(param)
                
        if 'removed' in params_diff:
            for param in params_diff['removed']:
                p = self.doc.add_paragraph(style='List Bullet 3')
                self._add_badge(p, "REMOVED PARAM", "DC3545", "FFFFFF")
                p.add_run(param)
                
        if 'modified' in params_diff:
            for param, changes in params_diff['modified'].items():
                p = self.doc.add_paragraph(style='List Bullet 3')
                self._add_badge(p, "MODIFIED PARAM", "FFC107", "000000")
                p.add_run(f"{param}")
                
                # Check for schema changes specifically
                if 'schema' in changes:
                    self._render_schema_diff_details(changes['schema'])
                
                # Other simple changes
                for key, val in changes.items():
                    if key == 'schema': continue
                    if isinstance(val, dict) and 'old' in val and 'new' in val:
                        p.add_run(f", {key} changed from '{val['old']}' to '{val['new']}'")

    def _add_request_body_changes(self, rb_diff):
        p = self.doc.add_paragraph(style='List Bullet 3')
        self._add_badge(p, "REQUEST BODY", "17A2B8", "FFFFFF")
        
        if 'required' in rb_diff:
            val = rb_diff['required']
            p.add_run(f"Required changed from {val['old']} to {val['new']}")
            
        if 'content' in rb_diff:
            content_diff = rb_diff['content']
            if 'modified' in content_diff:
                for media_type, mt_changes in content_diff['modified'].items():
                    if 'schema' in mt_changes:
                        p_mt = self.doc.add_paragraph(style='List Bullet 4')
                        p_mt.add_run(f"Media Type: {media_type}")
                        self._render_schema_diff_details(mt_changes['schema'])

    def _add_response_changes(self, responses_diff):
        if 'new' in responses_diff:
            for code in responses_diff['new']:
                p = self.doc.add_paragraph(style='List Bullet 3')
                self._add_badge(p, "NEW RESPONSE", "28A745", "FFFFFF")
                p.add_run(str(code))
                
        if 'removed' in responses_diff:
            for code in responses_diff['removed']:
                p = self.doc.add_paragraph(style='List Bullet 3')
                self._add_badge(p, "REMOVED RESPONSE", "DC3545", "FFFFFF")
                p.add_run(str(code))
                
        if 'modified' in responses_diff:
            for code, changes in responses_diff['modified'].items():
                p = self.doc.add_paragraph(style='List Bullet 3')
                self._add_badge(p, "MODIFIED RESPONSE", "FFC107", "000000")
                p.add_run(f"{code}")
                
                if 'description' in changes:
                    val = changes['description']
                    p.add_run(f" - Description changed")
                
                if 'content' in changes:
                    content_diff = changes['content']
                    if 'modified' in content_diff:
                        for media_type, mt_changes in content_diff['modified'].items():
                            if 'schema' in mt_changes:
                                p_mt = self.doc.add_paragraph(style='List Bullet 4')
                                p_mt.add_run(f"Media Type: {media_type}")
                                self._render_schema_diff_details(mt_changes['schema'])

    def _render_schema_diff_details(self, changes):
        """Renders detailed schema changes (tables, badges) inline."""
        
        # 1. $ref Change
        if '$ref' in changes:
            val = changes['$ref']
            p = self.doc.add_paragraph(style='List Bullet 4')
            self._add_badge(p, "SCHEMA REF CHANGED", "FFC107", "000000")
            p.add_run(f"From '{val['old']}' to '{val['new']}'")
            return # Usually if ref changes, we don't show internal diffs unless we deep compared them? 
                   # Comparator currently stops at ref change.

        # 2. Properties Table
        if 'properties' in changes and 'modified' in changes['properties']:
            self.doc.add_paragraph('Inline Schema Property Changes:', style='List Bullet 4')
            widths = [Inches(1.5), Inches(0.8), Inches(2.35), Inches(2.35)]
            
            # Indent table
            table = self._create_table(4, widths)
            tblPr = table._tblPr
            tblInd = get_or_add_child(tblPr, 'w:tblInd', TBL_PR_ORDER)
            tblInd.set(qn('w:w'), str(int(Inches(0.5).twips))) # Deeper indent for inline
            tblInd.set(qn('w:type'), 'dxa')
            
            self._style_header_row(table.rows[0], ['Property', 'Change', 'Old', 'New'])
            
            for prop, p_diff in changes['properties']['modified'].items():
                row = table.add_row()
                row.cells[0].text = prop
                row.cells[1].text = 'Mod'
                
                old_lines = []
                new_lines = []
                for k, v in p_diff.items():
                    old_lines.append(f"{k}: {v['old']}")
                    new_lines.append(f"{k}: {v['new']}")
                
                row.cells[2].text = '\n'.join(old_lines)
                row.cells[3].text = '\n'.join(new_lines)
                
                for cell in row.cells:
                    self._style_body_cell(cell)
            
            self.doc.add_paragraph().paragraph_format.space_after = Pt(6)

        # 3. Property Additions/Removals
        if 'properties' in changes:
            if 'new' in changes['properties'] and changes['properties']['new']:
                p = self.doc.add_paragraph(style='List Bullet 4')
                self._add_badge(p, "NEW PROP", "28A745", "FFFFFF")
                p.add_run(f"{', '.join(changes['properties']['new'])}")
                
            if 'removed' in changes['properties'] and changes['properties']['removed']:
                p = self.doc.add_paragraph(style='List Bullet 4')
                self._add_badge(p, "REMOVED PROP", "DC3545", "FFFFFF")
                p.add_run(f"{', '.join(changes['properties']['removed'])}")

        # 4. Other Schema Changes (oneOf, enum, etc.)
        for key, val in changes.items():
            if key in ['properties', '$ref']:
                continue
                
            p = self.doc.add_paragraph(style='List Bullet 4')
            self._add_badge(p, key.upper(), "17A2B8", "FFFFFF")
            
            if isinstance(val, dict) and ('added' in val or 'removed' in val):
                 if 'added' in val and val['added']:
                    p.add_run("Added options:")
                    for item in val['added']:
                        p_sub = self.doc.add_paragraph(style='List Bullet 5')
                        self._add_badge(p_sub, "ADDED", "28A745", "FFFFFF")
                        p_sub.add_run(self._format_schema_summary(item))
                 if 'removed' in val and val['removed']:
                    if 'added' in val and val['added']:
                        p = self.doc.add_paragraph(style='List Bullet 4')
                        self._add_badge(p, key.upper(), "17A2B8", "FFFFFF")
                    p.add_run("Removed options:")
                    for item in val['removed']:
                        p_sub = self.doc.add_paragraph(style='List Bullet 5')
                        self._add_badge(p_sub, "REMOVED", "DC3545", "FFFFFF")
                        p_sub.add_run(self._format_schema_summary(item))

            elif isinstance(val, dict) and 'old' in val and 'new' in val:
                p.add_run(f"Changed from '{val['old']}' to '{val['new']}'")
            elif isinstance(val, dict) and 'old_count' in val:
                p.add_run(f"Count changed from {val['old_count']} to {val['new_count']}")
            else:
                p.add_run(str(val))

    def _add_components(self):
        self.doc.add_heading('Components', 1)
        
        # Schemas
        schemas = self.diff.modified_components.get('schemas', {})
        if schemas:
            self.doc.add_heading('Schemas', 2)
            for schema, changes in schemas.items():
                # Logic Fix: Skip if no actual changes detected (empty dicts)
                if not changes: 
                    continue
                    
                self.doc.add_heading(schema, 3)
                
                # Properties Table
                if 'properties' in changes and 'modified' in changes['properties']:
                    self.doc.add_paragraph('Property Changes:')
                    widths = [Inches(1.5), Inches(0.8), Inches(2.35), Inches(2.35)]
                    
                    # Indent table
                    table = self._create_table(4, widths)
                    tblPr = table._tblPr
                    tblInd = get_or_add_child(tblPr, 'w:tblInd', TBL_PR_ORDER)
                    tblInd.set(qn('w:w'), str(int(Inches(0.25).twips))) # Indent 0.25 inch
                    tblInd.set(qn('w:type'), 'dxa')
                    
                    self._style_header_row(table.rows[0], ['Property', 'Change', 'Old', 'New'])
                    
                    for prop, p_diff in changes['properties']['modified'].items():
                        row = table.add_row()
                        row.cells[0].text = prop
                        row.cells[1].text = 'Mod'
                        
                        # Build multi-line text for old/new
                        old_lines = []
                        new_lines = []
                        for k, v in p_diff.items():
                            old_lines.append(f"{k}: {v['old']}")
                            new_lines.append(f"{k}: {v['new']}")
                        
                        row.cells[2].text = '\n'.join(old_lines)
                        row.cells[3].text = '\n'.join(new_lines)
                        
                        for cell in row.cells:
                            self._style_body_cell(cell)
                    
                    # Add spacing after table
                    self.doc.add_paragraph().paragraph_format.space_after = Pt(12)
                
                # Property Additions/Removals
                if 'properties' in changes:
                    if 'new' in changes['properties'] and changes['properties']['new']:
                        p = self.doc.add_paragraph(style='List Bullet')
                        self._add_badge(p, "NEW PROP", "28A745", "FFFFFF")
                        p.add_run(f"{', '.join(changes['properties']['new'])}")
                        
                    if 'removed' in changes['properties'] and changes['properties']['removed']:
                        p = self.doc.add_paragraph(style='List Bullet')
                        self._add_badge(p, "REMOVED PROP", "DC3545", "FFFFFF")
                        p.add_run(f"{', '.join(changes['properties']['removed'])}")

                # Other Schema Changes (oneOf, enum, required, etc.)
                for key, val in changes.items():
                    if key == 'properties':
                        continue
                        
                    p = self.doc.add_paragraph(style='List Bullet')
                    self._add_badge(p, key.upper(), "17A2B8", "FFFFFF") # Cyan badge for structural changes
                    
                    if isinstance(val, dict) and ('added' in val or 'removed' in val):
                        # Handle detailed combinator changes
                        if 'added' in val and val['added']:
                            p.add_run("Added options:")
                            for item in val['added']:
                                p_sub = self.doc.add_paragraph(style='List Bullet 2')
                                self._add_badge(p_sub, "ADDED", "28A745", "FFFFFF")
                                p_sub.add_run(self._format_schema_summary(item))
                                
                        if 'removed' in val and val['removed']:
                            # If we already added text to p, start a new one, else use p
                            if 'added' in val and val['added']:
                                p = self.doc.add_paragraph(style='List Bullet')
                                self._add_badge(p, key.upper(), "17A2B8", "FFFFFF")
                                
                            p.add_run("Removed options:")
                            for item in val['removed']:
                                p_sub = self.doc.add_paragraph(style='List Bullet 2')
                                self._add_badge(p_sub, "REMOVED", "DC3545", "FFFFFF")
                                p_sub.add_run(self._format_schema_summary(item))

                    elif isinstance(val, dict) and 'old' in val and 'new' in val:
                        p.add_run(f"Changed from '{val['old']}' to '{val['new']}'")
                    elif isinstance(val, dict) and 'old_count' in val:
                        p.add_run(f"Count changed from {val['old_count']} to {val['new_count']}")
                    else:
                        p.add_run(str(val))
        else:
            self.doc.add_paragraph('No modified schemas.')

    def _format_schema_summary(self, schema):
        if '$ref' in schema:
            return schema['$ref']
        elif 'type' in schema:
            return f"Type: {schema['type']}"
        return "Complex Schema Object"
