from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

# Helper to insert OXML elements in correct order
def get_or_add_child(parent, tag_name, order_list):
    # Check if exists
    child = parent.find(qn(tag_name))
    if child is not None:
        return child
    
    # Create new
    child = OxmlElement(tag_name)
    
    # Find insertion point
    # We want to insert BEFORE the first element that appears AFTER our tag in the order_list
    try:
        my_idx = order_list.index(tag_name)
    except ValueError:
        # If tag not in order list, append to end (risky but fallback)
        parent.append(child)
        return child
        
    for i in range(my_idx + 1, len(order_list)):
        next_tag = order_list[i]
        next_element = parent.find(qn(next_tag))
        if next_element is not None:
            parent.insert(parent.index(next_element), child)
            return child
            
    # If no later element found, append
    parent.append(child)
    return child

# ECMA-376 tblPr child order (simplified for common elements)
TBL_PR_ORDER = [
    'w:tblStyle', 'w:tblpPr', 'w:tblOverlap', 'w:bidiVisual',
    'w:tblStyleRowBandSize', 'w:tblStyleColBandSize', 
    'w:tblW', # Width
    'w:jc', 'w:tblCellSpacing', 'w:tblInd', 
    'w:tblBorders', 'w:shd', 
    'w:tblLayout', # Layout
    'w:tblCellMar', 'w:tblLook'
]

# ECMA-376 pPr child order (simplified)
P_PR_ORDER = [
    'w:pStyle', 'w:keepNext', 'w:keepLines', 'w:pageBreakBefore', 'w:framePr', 'w:widowControl',
    'w:numPr', 'w:suppressLineNumbers', 
    'w:pBdr', # Border
    'w:shd', 'w:tabs', 'w:suppressAutoHyphens', 'w:kinsoku', 'w:wordWrap', 'w:overflowPunct', 'w:topLinePunct', 'w:autoSpaceDE', 'w:autoSpaceDN', 'w:bidi', 'w:adjustRightInd', 'w:snapToGrid', 
    'w:spacing', # Spacing
    'w:ind', 'w:contextualSpacing', 'w:mirrorIndents', 'w:suppressOverlap', 'w:jc', 'w:textDirection', 'w:textAlignment', 'w:textboxTightWrap', 'w:outlineLvl', 'w:divId', 'w:cnfStyle', 'w:rPr', 'w:sectPr', 'w:pPrChange'
]

# ECMA-376 tcPr child order (simplified)
TC_PR_ORDER = [
    'w:cnfStyle', 'w:tcW', 'w:gridSpan', 'w:hMerge', 'w:vMerge', 
    'w:tcBorders', 
    'w:shd', # Shading
    'w:noWrap', 
    'w:tcMar', # Margins
    'w:textDirection', 'w:tcFitText', 'w:vAlign', 'w:hideMark', 'w:headers', 'w:cellIns', 'w:cellDel', 'w:cellMerge', 'w:tcPrChange'
]

def set_style(doc):
    # Modify Normal style (for body text)
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Segoe UI'
    font.size = Pt(10)
    style.paragraph_format.space_after = Pt(8)
    style.paragraph_format.line_spacing = 1.15

    # Create a specific style for Table content (compact)
    styles = doc.styles
    if 'Table Text' not in styles:
        table_style = styles.add_style('Table Text', 1) # 1 = Paragraph style
        table_font = table_style.font
        table_font.name = 'Segoe UI'
        table_font.size = Pt(9) # Slightly smaller for data
        table_style.paragraph_format.space_after = Pt(2) # Minimal spacing
        table_style.paragraph_format.line_spacing = 1.0 # Single spacing

    # Modify Heading 1
    h1 = doc.styles['Heading 1']
    h1.font.name = 'Segoe UI'
    h1.font.size = Pt(16)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(0, 51, 102) # Dark Navy
    h1.paragraph_format.space_before = Pt(24)
    h1.paragraph_format.space_after = Pt(12)
    
    # Add a border to Heading 1 using safe insertion
    pPr = h1._element.get_or_add_pPr()
    pbdr = get_or_add_child(pPr, 'w:pBdr', P_PR_ORDER)
    pbdr.clear() # Clear existing
    
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'auto')
    pbdr.append(bottom)

    # Modify Heading 2
    h2 = doc.styles['Heading 2']
    h2.font.name = 'Segoe UI'
    h2.font.size = Pt(13)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor(0, 51, 102)
    h2.paragraph_format.space_before = Pt(18)
    h2.paragraph_format.space_after = Pt(6)

    # Modify Heading 3
    h3 = doc.styles['Heading 3']
    h3.font.name = 'Segoe UI'
    h3.font.size = Pt(11)
    h3.font.bold = True
    h3.font.color.rgb = RGBColor(51, 51, 51) # Dark Grey

def setup_page_layout(doc):
    # Set narrow margins (0.5 inch) to maximize space
    section = doc.sections[0]
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)

def add_header_footer(doc):
    section = doc.sections[0]
    
    # Header
    header = section.header
    paragraph = header.paragraphs[0]
    paragraph.text = "OpenAPI Comparison Report - Confidential"
    paragraph.style = doc.styles['Normal']
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.runs[0]
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(128, 128, 128)

    # Footer
    footer = section.footer
    paragraph = footer.paragraphs[0]
    tab_stops = paragraph.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Inches(7.5), WD_TAB_ALIGNMENT.RIGHT)
    
    paragraph.text = f"Generated on {datetime.datetime.now().strftime('%Y-%m-%d')}\tPage "
    paragraph.style = doc.styles['Normal']
    run = paragraph.runs[0]
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(128, 128, 128)
    
    # Add page number field
    run_begin = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    run_begin._r.append(fldChar1)
    
    run_instr = paragraph.add_run()
    instrText = OxmlElement('w:instrText')
    instrText.text = "PAGE"
    run_instr._r.append(instrText)
    
    run_end = paragraph.add_run()
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run_end._r.append(fldChar2)

def set_cell_margins(cell, top=50, start=100, bottom=50, end=100):
    # Add padding to cell (values in twips, 1/20 of a point)
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = get_or_add_child(tcPr, 'w:tcMar', TC_PR_ORDER)
    tcMar.clear()
    
    for tag, val in [('w:top', top), ('w:start', start), ('w:bottom', bottom), ('w:end', end)]:
        node = OxmlElement(tag)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)

def create_table_header(table, headers):
    # Set header row to repeat on every page
    tr = table.rows[0]._tr
    trPr = tr.get_or_add_trPr()
    tblHeader = OxmlElement('w:tblHeader')
    trPr.append(tblHeader)

    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        paragraph = hdr_cells[i].paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = paragraph.runs[0]
        run.font.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(255, 255, 255)
        
        # Add shading
        tcPr = hdr_cells[i]._element.get_or_add_tcPr()
        shd = get_or_add_child(tcPr, 'w:shd', TC_PR_ORDER)
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '003366')
        
        # Add padding to header cells too
        set_cell_margins(hdr_cells[i], top=100, bottom=100)

def set_col_widths(table, widths):
    # Update table grid (tblGrid) to match these widths
    tblPr = table._tblPr
    tblGrid = table._element.find(qn('w:tblGrid'))
    if tblGrid is None:
        tblGrid = OxmlElement('w:tblGrid')
        # Insert grid after tblPr
        table._element.insert(table._element.index(tblPr) + 1, tblGrid)
    else:
        tblGrid.clear()
    
    for width in widths:
        gridCol = OxmlElement('w:gridCol')
        gridCol.set(qn('w:w'), str(int(width.twips)))
        tblGrid.append(gridCol)

    for row in table.rows:
        for idx, width in enumerate(widths):
            if idx < len(row.cells):
                cell = row.cells[idx]
                cell.width = width
                set_cell_margins(cell)
                for paragraph in cell.paragraphs:
                    paragraph.style = 'Table Text'

def set_fixed_layout(table):
    tblPr = table._tblPr
    tblLayout = get_or_add_child(tblPr, 'w:tblLayout', TBL_PR_ORDER)
    tblLayout.set(qn('w:type'), 'fixed')

def set_table_width(table, width):
    tblPr = table._tblPr
    tblW = get_or_add_child(tblPr, 'w:tblW', TBL_PR_ORDER)
    tblW.set(qn('w:w'), str(int(width.twips)))
    tblW.set(qn('w:type'), 'dxa')

def create_synthetic_template():
    doc = Document()
    set_style(doc)
    setup_page_layout(doc)
    add_header_footer(doc)
    
    title = doc.add_heading('OpenAPI Changes Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('Executive Summary', style='Subtitle').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    # Dashboard
    doc.add_heading('Dashboard', 1)
    table = doc.add_table(rows=2, cols=3)
    table.style = 'Table Grid'
    table.autofit = False 
    set_table_width(table, Inches(7.0))
    set_fixed_layout(table)
    set_col_widths(table, [Inches(2.3), Inches(2.3), Inches(2.4)])
    create_table_header(table, ['New Endpoints', 'Removed Endpoints', 'Modified Schemas'])
    
    row = table.rows[1]
    row.cells[0].text = '{{ diff.new_paths | length }}'
    row.cells[1].text = '{{ diff.removed_paths | length }}'
    row.cells[2].text = '{{ diff.modified_components.get("schemas", {}) | length }}'
    
    for cell in row.cells:
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.paragraphs[0].style = 'Table Text'

    # General Info
    doc.add_heading('General Info', 1)
    doc.add_paragraph('{%- if diff.info_changes -%}')
    
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    table.autofit = False
    set_table_width(table, Inches(7.0))
    set_fixed_layout(table)
    set_col_widths(table, [Inches(1.5), Inches(2.75), Inches(2.75)])
    create_table_header(table, ['Field', 'Old Value', 'New Value'])
    
    row_cells = table.add_row().cells
    row_cells[0].text = '{%- for key, change in diff.info_changes.items() -%}{{ key }}'
    row_cells[1].text = '{{ change.old }}'
    row_cells[2].text = '{{ change.new }}{%- endfor -%}'
    
    # Apply style to the template row
    for cell in row_cells:
        for p in cell.paragraphs:
            p.style = 'Table Text'
    
    doc.add_paragraph('{%- else -%}')
    doc.add_paragraph('No changes in General Info.')
    doc.add_paragraph('{%- endif -%}')
    
    # Endpoints
    doc.add_heading('Endpoints Summary', 1)
    
    doc.add_heading('New Endpoints', 2)
    doc.add_paragraph('{%- if diff.new_paths -%}')
    doc.add_paragraph('{%- for path in diff.new_paths -%}')
    doc.add_paragraph('{{ path }}', style='List Bullet')
    doc.add_paragraph('{%- endfor -%}')
    doc.add_paragraph('{%- else -%}')
    doc.add_paragraph('No new endpoints.')
    doc.add_paragraph('{%- endif -%}')
    
    doc.add_heading('Modified Endpoints', 2)
    doc.add_paragraph('{%- if diff.modified_paths -%}')
    doc.add_paragraph('{%- for path, changes in diff.modified_paths.items() -%}')
    doc.add_paragraph('{{ path }}', style='List Bullet')
    doc.add_paragraph('{%- if changes.new_ops -%}')
    doc.add_paragraph('New Operations: {{ changes.new_ops | join(", ") }}', style='List Bullet 2')
    doc.add_paragraph('{%- endif -%}')
    doc.add_paragraph('{%- endfor -%}')
    doc.add_paragraph('{%- else -%}')
    doc.add_paragraph('No modified endpoints.')
    doc.add_paragraph('{%- endif -%}')

    doc.save('templates/synthetic.docx')

def create_verbose_template():
    doc = Document()
    set_style(doc)
    setup_page_layout(doc)
    add_header_footer(doc)
    
    title = doc.add_heading('OpenAPI Changes Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('Detailed Developer Report', style='Subtitle').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    
    # Dashboard
    doc.add_heading('Dashboard', 1)
    table = doc.add_table(rows=2, cols=3)
    table.style = 'Table Grid'
    table.autofit = False
    set_table_width(table, Inches(7.0))
    set_fixed_layout(table)
    set_col_widths(table, [Inches(2.3), Inches(2.3), Inches(2.4)])
    create_table_header(table, ['New Endpoints', 'Removed Endpoints', 'Modified Schemas'])
    
    row = table.rows[1]
    row.cells[0].text = '{{ diff.new_paths | length }}'
    row.cells[1].text = '{{ diff.removed_paths | length }}'
    row.cells[2].text = '{{ diff.modified_components.get("schemas", {}) | length }}'
    for cell in row.cells:
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.paragraphs[0].style = 'Table Text'

    # General Info
    doc.add_heading('General Info', 1)
    doc.add_paragraph('{%- if diff.info_changes -%}')
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    table.autofit = False
    set_table_width(table, Inches(7.0))
    set_fixed_layout(table)
    set_col_widths(table, [Inches(1.5), Inches(2.75), Inches(2.75)])
    create_table_header(table, ['Field', 'Old Value', 'New Value'])
    
    row_cells = table.add_row().cells
    row_cells[0].text = '{%- for key, change in diff.info_changes.items() -%}{{ key }}'
    row_cells[1].text = '{{ change.old }}'
    row_cells[2].text = '{{ change.new }}{%- endfor -%}'
    
    for cell in row_cells:
        for p in cell.paragraphs:
            p.style = 'Table Text'
    
    doc.add_paragraph('{%- else -%}')
    doc.add_paragraph('No changes in General Info.')
    doc.add_paragraph('{%- endif -%}')

    # Paths
    doc.add_heading('Paths', 1)
    
    doc.add_heading('New Paths', 2)
    doc.add_paragraph('{%- if diff.new_paths -%}')
    doc.add_paragraph('{%- for path in diff.new_paths -%}')
    doc.add_paragraph('{{ path }}', style='List Bullet')
    doc.add_paragraph('{%- endfor -%}')
    doc.add_paragraph('{%- else -%}')
    doc.add_paragraph('None')
    doc.add_paragraph('{%- endif -%}')
    
    doc.add_heading('Modified Paths', 2)
    doc.add_paragraph('{%- if diff.modified_paths -%}')
    doc.add_paragraph('{%- for path, changes in diff.modified_paths.items() -%}')
    doc.add_heading('{{ path }}', 3)
    
    doc.add_paragraph('{%- if changes.new_ops -%}')
    doc.add_paragraph('New Operations: {{ changes.new_ops | join(", ") }}', style='List Bullet')
    doc.add_paragraph('{%- endif -%}')
    
    doc.add_paragraph('{%- if changes.modified_ops -%}')
    doc.add_paragraph('Modified Operations:', style='List Bullet')
    doc.add_paragraph('{%- for op, op_diff in changes.modified_ops.items() -%}')
    doc.add_paragraph('{{ op.upper() }}', style='List Bullet 2')
    doc.add_paragraph('{%- endfor -%}')
    doc.add_paragraph('{%- endif -%}')
    
    doc.add_paragraph('{%- endfor -%}')
    doc.add_paragraph('{%- else -%}')
    doc.add_paragraph('None')
    doc.add_paragraph('{%- endif -%}')
    
    # Components
    doc.add_heading('Components', 1)
    doc.add_heading('Modified Schemas', 2)
    doc.add_paragraph('{%- if diff.modified_components.get("schemas") -%}')
    doc.add_paragraph('{%- for schema, changes in diff.modified_components.schemas.items() -%}')
    doc.add_heading('{{ schema }}', 3)
    
    # Property Comparison Table Logic
    doc.add_paragraph('{%- for key, val in changes.items() -%}')
    doc.add_paragraph('{%- if key == "properties" and val.modified -%}')
    
    doc.add_paragraph('Property Changes:', style='Normal')
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    table.autofit = False
    set_table_width(table, Inches(7.0))
    set_fixed_layout(table)
    set_col_widths(table, [Inches(1.5), Inches(0.8), Inches(2.35), Inches(2.35)])
    create_table_header(table, ['Property', 'Change Type', 'Old Value', 'New Value'])
    
    row_cells = table.add_row().cells
    # Loop starts in first cell
    row_cells[0].text = '{%- for prop, p_diff in val.modified.items() -%}{{ prop }}'
    row_cells[1].text = 'Modified'
    # Inner loop for old values
    row_cells[2].text = '{%- for p_key, p_val in p_diff.items() -%}{{ p_key }}: {{ p_val.old }}{%- if not loop.last -%}\n{%- endif -%}{%- endfor -%}'
    # Inner loop for new values AND closing outer loop
    row_cells[3].text = '{%- for p_key, p_val in p_diff.items() -%}{{ p_key }}: {{ p_val.new }}{%- if not loop.last -%}\n{%- endif -%}{%- endfor -%}{%- endfor -%}'
    
    for cell in row_cells:
        for p in cell.paragraphs:
            p.style = 'Table Text'
    
    doc.add_paragraph('{%- endif -%}')
    
    # Handle New Properties
    doc.add_paragraph('{%- if key == "properties" and val.new -%}')
    doc.add_paragraph('New Properties: {{ val.new | join(", ") }}', style='List Bullet')
    doc.add_paragraph('{%- endif -%}')
    
    # Handle Removed Properties
    doc.add_paragraph('{%- if key == "properties" and val.removed -%}')
    doc.add_paragraph('Removed Properties: {{ val.removed | join(", ") }}', style='List Bullet')
    doc.add_paragraph('{%- endif -%}')
    
    # Handle Non-property changes (e.g. type, required)
    doc.add_paragraph('{%- if key != "properties" -%}')
    doc.add_paragraph('{{ key }}: {{ val.old }} -> {{ val.new }}', style='List Bullet')
    doc.add_paragraph('{%- endif -%}')
    
    doc.add_paragraph('{%- endfor -%}') # End changes loop
    doc.add_paragraph('{%- endfor -%}') # End schemas loop
    
    doc.add_paragraph('{%- else -%}')
    doc.add_paragraph('None')
    doc.add_paragraph('{%- endif -%}')

    doc.save('templates/verbose.docx')

if __name__ == '__main__':
    create_synthetic_template()
    create_verbose_template()
