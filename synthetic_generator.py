from analytic_generator import AnalyticDocxGenerator, get_or_add_child
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

class SyntheticDocxGenerator(AnalyticDocxGenerator):
    def generate(self, output_path):
        # Always add title to body
        self.doc.add_heading('OpenAPI Comparison - Synthesis Report', 0)
        
        self._add_spec_metadata()
        self._add_dashboard() # Change Matrix
        self._add_general_info_synthetic()
        self._add_endpoints_synthetic()
        self._add_schemas_synthetic()
        
        # Variable Substitution (Final Step)
        self._process_template_variables()
        
        self.doc.save(output_path)
        print(f"Synthetic Report generated at {output_path}")

    def _add_general_info_synthetic(self):
        self.doc.add_heading('General Info', 1)
        if not self.diff.info_changes:
            self.doc.add_paragraph('No changes in General Info.')
            return

        # Use a table for cleaner look
        widths = [Inches(2.0), Inches(5.0)]
        table = self._create_table(2, widths)
        self._style_header_row(table.rows[0], ['Field', 'Change'])
        
        for key, val in self.diff.info_changes.items():
            row = table.add_row()
            
            # Field Name
            row.cells[0].text = str(key)
            self._style_body_cell(row.cells[0])
            row.cells[0].paragraphs[0].runs[0].font.bold = True
            
            # Change Value
            cell = row.cells[1]
            old_val = val.get('old')
            new_val = val.get('new')
            
            if old_val is None and new_val is not None:
                cell.text = "Added"
                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(40, 167, 69) # Green
            elif old_val is not None and new_val is None:
                cell.text = "Removed"
                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(220, 53, 69) # Red
            else:
                cell.text = "Modified"
                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 193, 7) # Yellow/Orange (Darker for text)
                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(133, 100, 4) # Darker Orange
            
            cell.paragraphs[0].runs[0].font.bold = True
            self._style_body_cell(cell)
            
        self.doc.add_paragraph().paragraph_format.space_after = Pt(12)

        self.doc.add_paragraph().paragraph_format.space_after = Pt(12)

        # Override to set correct report name in header (Only if no template)
        if not self.has_template:
            section = self.doc.sections[0]
            # Header
            h_para = section.header.paragraphs[0]
            h_para.text = "OpenAPI Synthetic Report"
            h_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            h_para.runs[0].font.size = Pt(8)
            h_para.runs[0].font.color.rgb = RGBColor(128, 128, 128)
            
            # Footer
            f_para = section.footer.paragraphs[0]
            f_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            import datetime
            f_para.text = f"Generated on {datetime.datetime.now().strftime('%Y-%m-%d')} | Page "
            f_para.runs[0].font.size = Pt(8)
            f_para.runs[0].font.color.rgb = RGBColor(128, 128, 128)
            self._add_page_number(f_para)

    def _add_endpoints_synthetic(self):
        self.doc.add_heading('Endpoints Summary', 1)
        
        # New Endpoints
        self.doc.add_heading('New Endpoints', 2)
        if self.diff.new_paths:
            for path in self.diff.new_paths:
                p = self.doc.add_paragraph(style='List Bullet')
                self._add_pill_badge(p, "NEW", "28A745")
                p.add_run(" " + path)
        else:
            self.doc.add_paragraph('No new endpoints.')

        # Removed Endpoints
        self.doc.add_heading('Removed Endpoints', 2)
        if self.diff.removed_paths:
            for path in self.diff.removed_paths:
                p = self.doc.add_paragraph(style='List Bullet')
                self._add_pill_badge(p, "REMOVED", "DC3545")
                p.add_run(" " + path)
        else:
            self.doc.add_paragraph('No removed endpoints.')

        # Modified Endpoints
        self.doc.add_heading('Modified Endpoints', 2)
        if self.diff.modified_paths:
            for path, changes in self.diff.modified_paths.items():
                self.doc.add_heading(path, 3)
                
                # New Operations
                if 'new_ops' in changes:
                    p = self.doc.add_paragraph(style='List Bullet')
                    self._add_pill_badge(p, "NEW OPS", "28A745")
                    p.add_run(f" {', '.join(changes['new_ops']).upper()}")

                # Removed Operations
                if 'removed_ops' in changes:
                    p = self.doc.add_paragraph(style='List Bullet')
                    self._add_pill_badge(p, "REMOVED OPS", "DC3545")
                    p.add_run(f" {', '.join(changes['removed_ops']).upper()}")

                # Modified Operations
                if 'modified_ops' in changes:
                    for op, op_changes in changes['modified_ops'].items():
                        p = self.doc.add_paragraph(style='List Bullet')
                        self._add_pill_badge(p, op.upper(), "FFC107")
                        
                        summary_parts = []
                        
                        # Metadata
                        meta_keys = [k for k in op_changes.keys() if k in ['summary', 'description', 'deprecated', 'operationId']]
                        if meta_keys:
                            summary_parts.append(f"Metadata ({', '.join(meta_keys)})")
                        
                        # Parameters
                        if 'parameters' in op_changes:
                            p_diff = op_changes['parameters']
                            p_actions = []
                            if 'new' in p_diff: p_actions.append(f"Added {len(p_diff['new'])}")
                            if 'removed' in p_diff: p_actions.append(f"Removed {len(p_diff['removed'])}")
                            if 'modified' in p_diff: p_actions.append(f"Modified {len(p_diff['modified'])}")
                            summary_parts.append(f"Parameters ({', '.join(p_actions)})")

                        # Request Body
                        if 'requestBody' in op_changes:
                            summary_parts.append("Request Body")

                        # Responses
                        if 'responses' in op_changes:
                            r_diff = op_changes['responses']
                            r_actions = []
                            if 'new' in r_diff: r_actions.append(f"Added {', '.join(map(str, r_diff['new']))}")
                            if 'removed' in r_diff: r_actions.append(f"Removed {', '.join(map(str, r_diff['removed']))}")
                            if 'modified' in r_diff: r_actions.append(f"Modified {', '.join(map(str, r_diff['modified']))}")
                            summary_parts.append(f"Responses ({', '.join(r_actions)})")
                            
                        p.add_run(f" {'; '.join(summary_parts)}")
        else:
            self.doc.add_paragraph('No modified endpoints.')

    def _add_schemas_synthetic(self):
        # Calculate Pure Renames
        renamed_map = self.diff.renamed_components.get('schemas', {})
        modified_map = self.diff.modified_components.get('schemas', {})
        
        pure_renames = {}
        
        for old, new in renamed_map.items():
            if old not in modified_map:
                pure_renames[old] = new
        
        # List Pure Renames
        if pure_renames:
            self.doc.add_heading('Renamed Schemas', 2)
            for old, new in pure_renames.items():
                p = self.doc.add_paragraph(style='List Bullet')
                self._add_pill_badge(p, "RENAMED", "17A2B8")
                # Mixed Bold Styling
                p.add_run(" " + new).bold = True
                p.add_run(f" (was {old})")

        # List New Schemas
        s_new_items = self.diff.new_components.get('schemas', [])
        if s_new_items:
            self.doc.add_heading('New Schemas', 2)
            for item in sorted(s_new_items):
                p = self.doc.add_paragraph(style='List Bullet')
                self._add_pill_badge(p, "NEW", "28A745")
                p.add_run(" " + item).bold = True

        # List Removed Schemas
        s_rem_items = self.diff.removed_components.get('schemas', [])
        if s_rem_items:
            self.doc.add_heading('Removed Schemas', 2)
            for item in sorted(s_rem_items):
                p = self.doc.add_paragraph(style='List Bullet')
                self._add_pill_badge(p, "REMOVED", "DC3545")
                p.add_run(" " + item).bold = True

        # List Modified Schemas
        # Filter: Show here ONLY if substantial modification
        filtered_mod_items = {}
        for item_name, changes in modified_map.items():
            if self._is_substantial_modification(item_name, changes, renamed_map):
                filtered_mod_items[item_name] = changes

        if filtered_mod_items:
            self.doc.add_heading('Modified Schemas', 2)
            for item_name in sorted(filtered_mod_items.keys()):
                # Check rename
                display_name = item_name
                rename_note = ""
                new_name = renamed_map.get(item_name)
                if new_name:
                    display_name = new_name
                    rename_note = f" (was {item_name})"

                p = self.doc.add_paragraph(style='List Bullet')
                if rename_note:
                    self._add_pill_badge(p, "RENAMED & MODIFIED", "FFC107")
                else:
                    self._add_pill_badge(p, "MODIFIED", "FFC107")
                
                p.add_run(" " + display_name).bold = True
                if rename_note:
                    p.add_run(rename_note)
